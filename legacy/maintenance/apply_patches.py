#!/usr/bin/env python3
"""
apply_patches.py - va toan bo loi da audit vao JobOS.

Chay tu goc repo:
    python jobos_patch/apply_patches.py --check     # chi kiem tra, khong sua
    python jobos_patch/apply_patches.py --apply     # sua that, co backup

Nguyen tac:
  - Moi patch phai khop CHINH XAC MOT LAN. Khong khop hoac khop nhieu lan
    => dung toan bo, khong sua gi. Khong bao gio sua nua voi.
  - Backup .bak.<timestamp> truoc khi ghi.
  - Idempotent: chay lai lan hai se bao "da va roi", khong hong them.
"""

from __future__ import annotations

import argparse
import shutil
import sys
import time
from pathlib import Path
from typing import List, NamedTuple


class Patch(NamedTuple):
    pid: str
    rel_path: str
    old: str
    new: str
    why: str


SVC = "services/profile-ingestion"

PATCHES: List[Patch] = []


# ==========================================================================
# C1 + C2 : lo hong lop chong-bia trong evidence unit builder
# ==========================================================================

PATCHES.append(Patch(
    pid="C2",
    rel_path=f"{SVC}/build_structured_evidence_units_qwen_v2.py",
    why="infer_evidence_strength dung substring khong word-boundary; "
        "'lab' khop 'collaborate', 'project' khop 'projection'. "
        "Thay bang module chung co word boundary.",
    old='''def infer_evidence_strength(section_text: str) -> str:
    low = lower_norm(section_text)

    if any(x in low for x in ["tools to learn next", "learn next", "job-market tools", "common job-description", "high-priority framework to align next"]):
        return "job_market_target"
    if any(x in low for x in ["not directly used", "if not directly used", "familiar", "exposure", "materials", "tool list"]):
        return "material_exposure"
    if any(x in low for x in ["project", "used to", "was used", "in your project"]):
        return "project_use"
    if any(x in low for x in ["lab", "labs", "controlled lab", "lab environment"]):
        return "direct_lab_use"
    if any(x in low for x in ["should", "positioning", "say ", "do not overclaim"]):
        return "guidance_only"

    return "coursework_exposure"''',
    new='''def infer_evidence_strength(section_text: str) -> str:
    # PATCH C2: delegate sang module chung. Word-boundary + yeu cau dong tu
    # chu dong cho khang dinh manh. Xem services/common/jobos_safety.py
    return _safety.infer_evidence_strength(section_text)''',
))

PATCHES.append(Patch(
    pid="C1",
    rel_path=f"{SVC}/build_structured_evidence_units_qwen_v2.py",
    why="Guard chi ha cap direct_lab_use, BO SOT project_use. "
        "qwen tra project_use cho Metasploit thi khong gi chan lai. "
        "Them clamp: model chi duoc HA, khong duoc NANG.",
    old='''    # Deterministic downgrade rules.
    low = lower_norm(section_text)
    if any(x in low for x in ["learn next", "job-market tools", "if not directly used", "not directly used"]):
        evidence_strength = "job_market_target" if "learn next" in low or "job-market" in low else "material_exposure"
        if claim_type == "tool_experience":
            claim_type = "tool_exposure"

    if any(x in low for x in ["familiar", "exposure", "materials", "tool list"]):
        if evidence_strength == "direct_lab_use":
            evidence_strength = "material_exposure"
        if claim_type == "tool_experience":
            claim_type = "tool_exposure"

    # Normalize claim_type against evidence_strength.
    if evidence_strength in {"direct_lab_use", "project_use"} and claim_type == "tool_exposure":
        claim_type = "tool_experience"

    if evidence_strength in {"material_exposure", "guidance_only"} and claim_type == "tool_experience":
        claim_type = "tool_exposure"

    if evidence_strength == "job_market_target":
        claim_type = "job_market_target"''',
    new='''    # PATCH C1: chot chan ba lop.
    # 1. model chi duoc HA cap so voi suy luan deterministic
    evidence_strength = _safety.clamp_strength(evidence_strength, det_strength)

    # 2. ha cap deterministic, ap cho MOI strength manh (khong chi direct_lab_use),
    #    cong denylist theo ten tool (Metasploit, sqlmap, Nessus, ...)
    evidence_strength, claim_type = _safety.apply_downgrades(
        evidence_strength, claim_type, section_text, tool_name
    )

    # 3. clamp lan cuoi phong truong hop apply_downgrades bi sua sau nay
    evidence_strength = _safety.clamp_strength(evidence_strength, det_strength)''',
))


# ==========================================================================
# H3 : do_not_overclaim_rules bi nhiem danh tu tran / enum
# ==========================================================================

PATCHES.append(Patch(
    pid="H3a",
    rel_path=f"{SVC}/synthesize_structured_tool_workflow_assets_qwen_v1.py",
    why="must_not_claim (cau cam) va does_not_support_claims (danh tu) bi gop "
        "vao mot mang. ~60% phan tu la danh tu tran, co ca enum direct_lab_use. "
        "L6 verifier doc mang nay nen rule rac = chan khong duoc.",
    old='''    do_not = clean_list(raw.get("do_not_overclaim_rules"))
    for u in units:
        do_not.extend(clean_list(u.get("must_not_claim")))
        do_not.extend(clean_list(u.get("does_not_support_claims")))
    do_not.append("Do not present academic lab, coursework, or project exposure as professional employment.")
    do_not.append("Do not claim production deployment, certification, or expert mastery unless separately supported.")
    do_not = list(dict.fromkeys(do_not))''',
    new='''    # PATCH H3: moi phan tu phai la CAU CAM. Danh tu tran duoc nang thanh
    # "Do not claim X."; enum/tag ky thuat (direct_lab_use, pki_tls) bi loai.
    do_not = _safety.build_overclaim_rules(
        raw.get("do_not_overclaim_rules"),
        *[u.get("must_not_claim") for u in units],
        *[u.get("does_not_support_claims") for u in units],
    )''',
))

PATCHES.append(Patch(
    pid="H3b",
    rel_path=f"{SVC}/synthesize_profile_assets_qwen_v1.py",
    why="Cung loi H3, lap lai o synthesizer thu hai. clean_list khong validate "
        "hinh thai nen danh tu tran lot thang vao rule.",
    old='''        "do_not_overclaim_rules": clean_list(asset.get("do_not_overclaim_rules")),''',
    new='''        "do_not_overclaim_rules": _safety.build_overclaim_rules(asset.get("do_not_overclaim_rules")),''',
))


# ==========================================================================
# H1 : soft-fail import -> mat du lieu im lang
# ==========================================================================

PATCHES.append(Patch(
    pid="H1",
    rel_path=f"{SVC}/ingest_files.py",
    why="import mem: thieu pypdf/python-docx thi bien = None, roi parse tra "
        "chuoi RONG ma khong exception. Moi file .docx/.pdf ghi vao DB rong, "
        "exit code 0. Day la data corruption im lang o stage dau tien.",
    old='''try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import docx
except Exception:
    docx = None''',
    new='''# PATCH H1: hard-fail. Thieu dependency thi chet ngay, khong parse rong im lang.
try:
    from pypdf import PdfReader
except ImportError as e:
    raise SystemExit(
        "FATAL: thieu 'pypdf'. Chay: pip install -r requirements.txt\\n"
        f"       chi tiet: {e}"
    )

try:
    import docx
except ImportError as e:
    raise SystemExit(
        "FATAL: thieu 'python-docx'. Chay: pip install -r requirements.txt\\n"
        f"       chi tiet: {e}"
    )''',
))

PATCHES.append(Patch(
    pid="H1b",
    rel_path=f"{SVC}/ingest_files.py",
    why="Bo nhanh chet 'pypdf_missing' - khong con toi duoc sau H1.",
    old='''    if PdfReader is None:
        return "", "pypdf_missing", "pypdf is not installed"

''',
    new="",
))

PATCHES.append(Patch(
    pid="H1c",
    rel_path=f"{SVC}/ingest_files.py",
    why="Bo nhanh chet 'python-docx_missing' - khong con toi duoc sau H1.",
    old='''    if docx is None:
        return "", "python-docx_missing", "python-docx is not installed"

''',
    new="",
))


# ==========================================================================
# H2 : parse_docx vut bo heading -> root cause cua chunk hong
# ==========================================================================

PATCHES.append(Patch(
    pid="H2",
    rel_path=f"{SVC}/ingest_files.py",
    why="para.style.name khong bao gio duoc doc, nen moi phan cap Heading "
        "bien mat. Downstream phai doan lai bang regex tren text da mat cau truc. "
        "Do la ly do file PDF chunk TOT HON file DOCX cung noi dung.",
    old='''        d = docx.Document(str(path))
        parts = []

        for para in d.paragraphs:
            txt = para.text.strip()
            if txt:
                parts.append(txt)''',
    new='''        d = docx.Document(str(path))
        parts = []

        # PATCH H2: giu phan cap Heading tu Word.
        #   Heading N  -> "#### Title"  (MARKDOWN_HEADING_RE o section builder bat duoc)
        #   nhan in dam ket thuc bang ":" -> "@@FIELD@@ Nhan:" (KHONG phai heading)
        # Nho vay detect_heading khong con phai doan, va khong the nham
        # "Portfolio positioning:" thanh tieu de section.
        for para in d.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue

            style = (para.style.name or "").lower() if para.style is not None else ""

            if style.startswith("heading"):
                digits = "".join(c for c in style if c.isdigit())
                level = min(int(digits), 5) if digits else 1
                parts.append("")
                parts.append(f"{'#' * level} {txt}")
                parts.append("")
                continue

            if style in ("title", "subtitle"):
                parts.append("")
                parts.append(f"# {txt}")
                parts.append("")
                continue

            is_label = (
                len(txt) < 60
                and txt.endswith(":")
                and para.runs
                and bool(para.runs[0].bold)
            )
            if is_label:
                parts.append(f"@@FIELD@@ {txt}")
                continue

            parts.append(txt)''',
))


# ==========================================================================
# H4 + M1 : VERSION cung + vut du lieu im lang o section builder
# ==========================================================================

PATCHES.append(Patch(
    pid="H4",
    rel_path=f"{SVC}/build_structured_sections_from_inventory_docs.py",
    why="VERSION la hang so cung, khong bump khi doi logic. DB dang chua "
        "output cua nhieu doi code cung mot nhan version, va "
        "existing_structured_count() chan re-run nen rac nam lai vinh vien.",
    old='''VERSION = "structured_section_boundary_v1_2026_04_27"''',
    new='''# PATCH H4: VERSION tu sinh theo hash cua logic. Doi logic -> doi version
# -> existing_structured_count() khong chan nua -> tu chay lai.
def _logic_version() -> str:
    import hashlib
    import inspect
    src = inspect.getsource(detect_heading) + inspect.getsource(looks_structured)
    return "structured_section_boundary_v2_" + hashlib.sha1(src.encode()).hexdigest()[:10]


VERSION = "PLACEHOLDER_SET_AT_BOTTOM"''',
))

PATCHES.append(Patch(
    pid="M1",
    rel_path=f"{SVC}/build_structured_sections_from_inventory_docs.py",
    why="Section bi vut (body<80, hoac looks_structured=False) khong duoc dem "
        "hay log. 'found: 81' khong cho biet bo bao nhieu. Khong audit duoc.",
    old='''    out: List[Tuple[str, str]] = []
    for title, body_lines in sections:
        body = "\\n".join(body_lines).strip()
        if len(body) < 80:
            continue
        if looks_structured(title, body):
            out.append((title[:500], body))

    return out''',
    new='''    # PATCH M1: dem va in so section bi vut, thay vi bo im lang.
    out: List[Tuple[str, str]] = []
    dropped_short = 0
    dropped_unstructured = 0
    for title, body_lines in sections:
        body = "\\n".join(body_lines).strip()
        if len(body) < 80:
            dropped_short += 1
            continue
        if not looks_structured(title, body):
            dropped_unstructured += 1
            continue
        out.append((title[:500], body))

    print(
        f"    split: raw={len(sections)} kept={len(out)} "
        f"dropped_short={dropped_short} dropped_unstructured={dropped_unstructured}"
    )
    return out''',
))


# ==========================================================================
# Import module chung - them vao dau moi file can dung
# ==========================================================================

IMPORT_BLOCK = '''
# PATCH: module an toan dung chung (evidence strength + overclaim rules).
import sys as _sys
from pathlib import Path as _Path

_COMMON = _Path(__file__).resolve().parents[1] / "common"
if str(_COMMON.parent) not in _sys.path:
    _sys.path.insert(0, str(_COMMON.parent))
from common import jobos_safety as _safety  # noqa: E402
'''

NEEDS_SAFETY_IMPORT = [
    f"{SVC}/build_structured_evidence_units_qwen_v2.py",
    f"{SVC}/synthesize_structured_tool_workflow_assets_qwen_v1.py",
    f"{SVC}/synthesize_profile_assets_qwen_v1.py",
]


# ==========================================================================
# Runner
# ==========================================================================

def find_repo_root(start: Path) -> Path:
    for p in [start, *start.parents]:
        if (p / SVC).is_dir():
            return p
    raise SystemExit(
        f"FATAL: khong tim thay '{SVC}/'. Chay script nay tu goc repo job-apply-os."
    )


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--apply", action="store_true", help="ghi that")
    ap.add_argument("--check", action="store_true", help="chi kiem tra")
    args = ap.parse_args()

    if not args.apply and not args.check:
        args.check = True

    root = find_repo_root(Path.cwd().resolve())
    print(f"repo root: {root}")
    print(f"mode:      {'APPLY' if args.apply else 'CHECK-ONLY'}")
    print("")

    stamp = time.strftime("%Y%m%d_%H%M%S")
    problems: List[str] = []
    plan: List[tuple] = []

    # --- 1. kiem tra tung patch ---
    for p in PATCHES:
        target = root / p.rel_path
        if not target.exists():
            problems.append(f"[{p.pid}] KHONG CO FILE: {p.rel_path}")
            continue
        src = target.read_text(encoding="utf-8")
        n = src.count(p.old)
        if n == 1:
            plan.append((p, target))
            print(f"[{p.pid}] OK       {p.rel_path}")
        elif n == 0 and (not p.new.strip() or p.new in src):
            # new rong = patch xoa; old bien mat nghia la da xoa roi.
            print(f"[{p.pid}] DA VA    {p.rel_path}  (bo qua)")
        elif n == 0:
            problems.append(
                f"[{p.pid}] KHONG KHOP trong {p.rel_path}\n"
                f"          -> file da bi sua tay. Doi chieu bang tay truoc khi va.\n"
                f"          -> ly do patch: {p.why}"
            )
        else:
            problems.append(f"[{p.pid}] KHOP {n} LAN trong {p.rel_path} (can dung 1)")

    # --- 2. import block ---
    import_plan: List[Path] = []
    for rel in NEEDS_SAFETY_IMPORT:
        target = root / rel
        if not target.exists():
            problems.append(f"[IMPORT] KHONG CO FILE: {rel}")
            continue
        src = target.read_text(encoding="utf-8")
        if "jobos_safety as _safety" in src:
            print(f"[IMPORT] DA CO   {rel}")
        else:
            import_plan.append(target)
            print(f"[IMPORT] OK      {rel}")

    # --- 3. module chung phai ton tai ---
    common_pkg = root / "services" / "common"
    if not (common_pkg / "jobos_safety.py").exists():
        problems.append(
            "[SETUP] thieu services/common/jobos_safety.py\n"
            "        -> copy tu jobos_patch/services/common/ truoc khi chay --apply"
        )

    print("")
    if problems:
        print("===== DUNG. Co van de: =====")
        for x in problems:
            print("  " + x)
        return 1

    if not args.apply:
        print(f"===== CHECK OK. {len(plan)} patch + {len(import_plan)} import san sang. =====")
        print("Chay lai voi --apply de ghi.")
        return 0

    # --- 4. ghi ---
    touched = set()
    for p, target in plan:
        if target not in touched:
            shutil.copy2(target, target.with_suffix(target.suffix + f".bak.{stamp}"))
            touched.add(target)
        src = target.read_text(encoding="utf-8")
        target.write_text(src.replace(p.old, p.new, 1), encoding="utf-8")
        print(f"[{p.pid}] da va  {p.rel_path}")

    for target in import_plan:
        if target not in touched:
            shutil.copy2(target, target.with_suffix(target.suffix + f".bak.{stamp}"))
            touched.add(target)
        src = target.read_text(encoding="utf-8")
        lines = src.split("\n")
        idx = 0
        for i, line in enumerate(lines):
            if line.startswith("import ") or line.startswith("from "):
                idx = i + 1
        lines.insert(idx, IMPORT_BLOCK)
        target.write_text("\n".join(lines), encoding="utf-8")
        print(f"[IMPORT] da them  {target.relative_to(root)}")

    # --- 5. H4 can dat VERSION o cuoi file (sau khi ham da dinh nghia) ---
    sec = root / SVC / "build_structured_sections_from_inventory_docs.py"
    if sec.exists():
        src = sec.read_text(encoding="utf-8")
        if 'VERSION = "PLACEHOLDER_SET_AT_BOTTOM"' in src:
            anchor = "def main() -> int:"
            if anchor in src:
                src = src.replace(anchor, "VERSION = _logic_version()\n\n\n" + anchor, 1)
                sec.write_text(src, encoding="utf-8")
                print("[H4] da dat VERSION = _logic_version() truoc main()")
            else:
                print("[H4] CANH BAO: khong tim thay 'def main()', dat VERSION bang tay!")

    print("")
    print(f"===== XONG. {len(touched)} file da sua. Backup: *.bak.{stamp} =====")
    print("Buoc tiep: python -m py_compile services/profile-ingestion/*.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
