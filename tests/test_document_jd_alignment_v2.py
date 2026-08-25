"""Regression coverage for JD-first resume/cover-letter tailoring V2."""
from __future__ import annotations

import importlib.util
import json
import os
from pathlib import Path
import sys
import types

from docx import Document

os.environ.setdefault("JOBOS_DB_PASSWORD", "unit-test")

if "psycopg" not in sys.modules:
    psycopg = types.ModuleType("psycopg")
    psycopg.connect = lambda *_a, **_k: None
    psycopg_types = types.ModuleType("psycopg.types")
    psycopg_json = types.ModuleType("psycopg.types.json")
    psycopg_json.Jsonb = lambda value: value
    sys.modules["psycopg"] = psycopg
    sys.modules["psycopg.types"] = psycopg_types
    sys.modules["psycopg.types.json"] = psycopg_json

ROOT = Path(__file__).resolve().parents[1]
GEN = ROOT / "services" / "document-generation" / "generate_documents_v1.py"
SPEC = importlib.util.spec_from_file_location("jobos_docgen_alignment_test", GEN)
docgen = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(docgen)

from services.common.document_prompt_templates_v1 import (
    build_cover_alignment_blueprint_prompt,
    build_cover_alignment_audit_prompt,
    build_cover_letter_tailoring_prompt,
    build_resume_tailoring_prompt,
    material_requirement_summary,
    requirement_catalog,
)
from services.common.resume_experience_bullet_audit import (
    load_template_experience_baselines,
    validate_experience_bullet_change,
)


def _app() -> dict:
    return {
        "company": "Example Co",
        "job_title": "Security Analyst",
        "role_family": "soc_dfir",
        "seniority_level": "entry",
        "jd_text": (
            "Investigate security alerts. Automate reporting with Python. "
            "Communicate clearly with stakeholders. Collaborate across teams. "
            "Learn new security tooling quickly. Protect customer systems."
        ),
        "matched_requirements": [
            {"requirement": "Investigate security alerts", "profile_support": "supported"},
            {"requirement": "Automate reporting with Python", "profile_support": "supported"},
            {"requirement": "Communicate clearly with stakeholders", "profile_support": "supported"},
            {"requirement": "Collaborate across teams", "profile_support": "supported"},
            {"requirement": "Learn new security tooling quickly", "profile_support": "supported"},
        ],
        "missing_or_weak_requirements": [{"requirement": "SIEM administration", "severity": "medium"}],
        "company_context": {},
    }


def test_resume_prompt_sees_full_jd_and_has_80_percent_grounded_coverage_policy():
    app = _app()
    prompt = build_resume_tailoring_prompt(
        app=app,
        asset_catalog="[ASSET a1] approved user evidence",
        max_project_bullets=8,
        fixed_projects=("1-2 Project",),
        fixed_project_asset_rules="- slots 1-2: a1",
        baseline_subtitles={1: "Old subtitle"},
        baseline_project_bullets={1: "Old bullet"},
        experience_baselines={1: {"header_context": "Security Intern | Example Co", "previous_bullet": "Old exp bullet", "max_chars": 80}},
    )
    assert app["jd_text"] in prompt
    assert "80%" in prompt
    assert "Experience bullet descriptions MAY be rewritten" in prompt
    assert "Keep every employer/company name, job title, date" in prompt
    assert "matched_requirement_ids" in prompt
    assert "NEVER invent" in prompt


def test_cover_prompts_use_full_jd_and_separate_positioning_from_technical_truth():
    app = _app()
    blueprint_prompt = build_cover_alignment_blueprint_prompt(app=app)
    assert app["jd_text"] in blueprint_prompt
    audit_prompt = build_cover_alignment_audit_prompt(
        app=app, alignment_blueprint={"soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}]}
    )
    assert app["jd_text"] in audit_prompt
    assert "NOT an exhaustive coverage gate" in audit_prompt
    assert "distinctive useful additions" in audit_prompt
    prompt = build_cover_letter_tailoring_prompt(
        app=app,
        asset_catalog="[ASSET a1] Python reporting automation",
        alignment_blueprint={
            "about_me_targets": [{"id": "A1", "quote": "Protect customer systems"}],
            "interest_targets": [{"id": "I1", "quote": "Protect customer systems"}],
            "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
            "technical_targets": [{"id": "T1", "quote": "Automate reporting with Python"}],
        },
    )
    assert app["jd_text"] in prompt
    assert "SELECTED FIT, NOT CHECKLIST COVERAGE" in prompt
    assert "2-5 total A/I/S targets" in prompt
    assert "technical_evidence" in prompt
    assert "approved user" in prompt.lower()
    assert "subjective positioning" in prompt.lower()


def test_requirement_catalog_assigns_stable_supportable_ids():
    items = requirement_catalog(_app()["matched_requirements"])
    assert [item["id"] for item in items] == ["M1", "M2", "M3", "M4", "M5"]


def test_experience_audit_binds_rewrite_to_immutable_job_header_and_user_evidence(tmp_path):
    template = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Security Intern | Example Co | Jan-May 2025")
    doc.add_paragraph("Investigated alerts and documented findings.", style="List Paragraph")
    doc.add_paragraph("PROJECTS")
    doc.save(template)
    baseline = load_template_experience_baselines(template)
    change = {
        "slot": 1,
        "header_context": "Security Intern | Example Co | Jan-May 2025",
        "previous_bullet": "Investigated alerts and documented findings.",
        "text": "Investigated security alerts and documented findings for analyst review.",
        "jd_requirement_quote": "Investigate security alerts",
        "experience_evidence_quote": "Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings",
        "word_change_rationale": [
            {"before": "alerts", "after": "security alerts", "why": "Uses the JD's precise security-alert wording without changing scope."},
            {"before": "documented findings", "after": "documented findings for analyst review", "why": "Keeps the evidenced documentation work while clarifying its review purpose."},
        ],
        "why_better": "Directly matches the alert-investigation JD language while preserving the original internship scope.",
    }
    assert validate_experience_bullet_change(
        change, baselines=baseline, jd_text=_app()["jd_text"],
        asset_source="Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings for analyst review.",
    ) == []
    bad = dict(change, header_context="Senior Security Engineer | Fake Co")
    assert any("header_context" in problem for problem in validate_experience_bullet_change(
        bad, baselines=baseline, jd_text=_app()["jd_text"],
        asset_source="Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings for analyst review.",
    ))


def _skill(category: str, item: str, req_id: str, jd_quote: str) -> dict:
    return {
        "category": category,
        "items": item,
        "source_asset_id": "asset-1",
        "matched_requirement_ids": [req_id],
        "jd_requirement_quote": jd_quote,
        "skill_evidence_quote": f"evidence for {item}",
    }


def test_resume_80_percent_is_non_blocking_quality_target_over_full_material_jd():
    app = _app()
    coverage = material_requirement_summary(app)
    assert coverage["total_material_requirements"] == 6
    assert coverage["truthful_coverage_ceiling_percent"] == 83.3
    assert coverage["required_supportable_covered_count"] == 5

    parsed = {
        "skill_lines_ranked": [
            _skill("Security", "Alert triage", "M1", "Investigate security alerts"),
            _skill("Automation", "Python", "M2", "Automate reporting with Python"),
            _skill("Communication", "Stakeholder communication", "M3", "Communicate clearly with stakeholders"),
            _skill("Collaboration", "Cross-team work", "M4", "Collaborate across teams"),
            _skill("Learning", "Security tooling", "M5", "Learn new security tooling quickly"),
        ],
        "not_supported": [], "self_check": "grounded",
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "resume", parsed, {"asset-1"}, jd_text=app["jd_text"],
        matched_requirement_ids={"M1", "M2", "M3", "M4", "M5"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=6,
    )
    assert content
    assert used == ["asset-1"]
    assert evidence["jd_alignment"]["coverage_percent"] == 83.3
    assert evidence["jd_alignment"]["truthful_coverage_ceiling_percent"] == 83.3
    assert evidence["jd_alignment"]["gate_passed"] is True

    # Four of five supportable requirements is only 66.7% of the six material
    # requirements. That misses the quality target but must not brick generation.
    parsed["skill_lines_ranked"] = parsed["skill_lines_ranked"][:4]
    content, _used, evidence, dropped = docgen.validate_and_render(
        "resume", parsed, {"asset-1"}, jd_text=app["jd_text"],
        matched_requirement_ids={"M1", "M2", "M3", "M4", "M5"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=6,
    )
    assert content
    assert evidence["jd_alignment"]["coverage_percent"] == 66.7
    assert evidence["jd_alignment"]["target_met"] is False
    assert evidence["jd_alignment"]["gate_passed"] is True
    assert any("below the 80% target" in item for item in evidence["warnings"])


def test_resume_when_80_percent_is_impossible_keeps_safe_lower_coverage_output():
    parsed = {
        "skill_lines_ranked": [
            _skill("Security", "Alert triage", "M1", "Investigate security alerts"),
            _skill("Automation", "Python", "M2", "Automate reporting with Python"),
            _skill("Communication", "Stakeholder communication", "M3", "Communicate clearly with stakeholders"),
        ],
        "not_supported": [], "self_check": "grounded",
    }
    content, _used, evidence, _dropped = docgen.validate_and_render(
        "resume", parsed, {"asset-1"}, jd_text=_app()["jd_text"],
        matched_requirement_ids={"M1", "M2", "M3"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=5,
    )
    assert content
    assert evidence["jd_alignment"]["coverage_percent"] == 60.0
    assert evidence["jd_alignment"]["truthful_coverage_ceiling_percent"] == 60.0
    assert evidence["jd_alignment"]["target_reachable_truthfully"] is False
    assert evidence["jd_alignment"]["gate_passed"] is True

    parsed["skill_lines_ranked"] = parsed["skill_lines_ranked"][:2]
    content, _used, evidence, _dropped = docgen.validate_and_render(
        "resume", parsed, {"asset-1"}, jd_text=_app()["jd_text"],
        matched_requirement_ids={"M1", "M2", "M3"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=5,
    )
    assert content
    assert evidence["jd_alignment"]["required_supportable_covered_count"] == 3
    assert evidence["jd_alignment"]["target_met"] is False
    assert evidence["jd_alignment"]["gate_passed"] is True
    assert evidence["warnings"]


def test_resume_generation_allows_jd_first_experience_rewrite_without_official_resume_evidence():
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 100,
        }
    }
    item = {
        "slot": 1,
        "header_context": baseline[1]["header_context"],
        "previous_bullet": baseline[1]["previous_bullet"],
        "text": "Investigated security alerts and documented findings for analyst review.",
        "source_asset_id": "none",
        "jd_requirement_quote": "Investigate security alerts",
        "experience_evidence_quote": "",
        "matched_requirement_ids": ["M1"],
        "word_change_rationale": [
            {"before": "alerts", "after": "security alerts", "why": "Matches the exact JD security-alert wording without changing the role."},
            {"before": "documented findings", "after": "documented findings for analyst review", "why": "Keeps a general analyst-review framing consistent with the existing job title."},
        ],
        "why_better": "Improves JD relevance while preserving the immutable internship title and general responsibility scope.",
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "resume", {"experience_updates": [item]}, set(),
        jd_text=_app()["jd_text"], experience_baselines=baseline,
        experience_source_asset_ids=set(), matched_requirement_ids={"M1"},
    )
    assert "Investigated security alerts" in content
    assert used == []
    assert evidence["resume_template"]["experience_bullets"][0]["source_asset_id"] is None
    assert not dropped

def test_cover_letter_uses_selected_ais_subset_and_keeps_technical_grounded():
    jd = _app()["jd_text"]
    blueprint = {
        "about_me_targets": [{"id": "A1", "quote": "Protect customer systems"}],
        "interest_targets": [{"id": "I1", "quote": "Protect customer systems"}],
        "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
        "technical_targets": [{"id": "T1", "quote": "Automate reporting with Python"}],
    }
    parsed = {
        "paragraphs": [
            {"sentences": [
                {"kind": "about_me_positioning", "text": "I am drawn to work centered on protecting customer systems.", "alignment_ids": ["A1"], "jd_requirement_quote": "Protect customer systems", "source_asset_id": "none"},
                {"kind": "role_interest", "text": "That customer-protection focus is the part of this role that interests me most.", "alignment_ids": ["I1"], "jd_requirement_quote": "Protect customer systems", "source_asset_id": "none"},
                {"kind": "soft_skill_positioning", "text": "I value collaborating across teams and keeping shared work clear.", "alignment_ids": ["S1"], "jd_requirement_quote": "Collaborate across teams", "source_asset_id": "none"},
            ]},
            {"sentences": [
                {"kind": "technical_evidence", "text": "I used Python to automate a recurring reporting workflow.", "alignment_ids": ["T1"], "jd_requirement_quote": "Automate reporting with Python", "source_asset_id": "asset-1", "candidate_evidence_quote": "automated recurring reporting with Python"},
            ]},
        ],
        "not_supported": [], "self_check": "selected A/I/S; technical grounded",
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "cover_letter", parsed, {"asset-1"}, jd_text=jd,
        company="Example Co", job_title="Security Analyst",
        cover_alignment_blueprint=blueprint,
    )
    assert content.startswith("I am applying for the Security Analyst position at Example Co.")
    assert used == ["asset-1"]
    assert evidence["cover_alignment"]["coverage_percent"] == 100.0
    assert evidence["cover_alignment"]["gate_passed"] is True
    assert {claim["kind"] for claim in evidence["claims"]} >= {
        "cover_letter_positioning", "cover_letter_evidence", "cover_letter_structure"
    }
    assert not dropped

    # Removing the soft-skill sentence is valid: the letter intentionally picks
    # only a few A/I/S themes rather than mirroring the whole JD.
    parsed["paragraphs"][0]["sentences"] = parsed["paragraphs"][0]["sentences"][:2]
    content, _used, evidence, dropped = docgen.validate_and_render(
        "cover_letter", parsed, {"asset-1"}, jd_text=jd,
        company="Example Co", job_title="Security Analyst",
        cover_alignment_blueprint=blueprint,
    )
    assert content
    assert evidence["cover_alignment"]["coverage_percent"] == 100.0
    assert evidence["cover_alignment"]["gate_passed"] is True
    assert "S1" in evidence["cover_alignment"]["available_positioning_ids"]
    assert "S1" not in evidence["cover_alignment"]["selected_positioning_ids"]
    assert not any("missing targets" in item for item in dropped)


def test_cover_blueprint_normalizer_drops_non_verbatim_targets():
    jd = _app()["jd_text"]
    normalized = docgen.normalize_cover_alignment_blueprint({
        "soft_skill_targets": [
            {"quote": "Collaborate across teams", "why": "explicit"},
            {"quote": "Be an amazing communicator", "why": "invented paraphrase"},
        ],
        "about_me_targets": [], "interest_targets": [], "technical_targets": [],
    }, jd)
    assert normalized["soft_skill_targets"] == [
        {"id": "S1", "quote": "Collaborate across teams", "why": "explicit"}
    ]


def test_cover_blueprint_candidate_audit_merge_adds_only_exact_jd_targets():
    jd = _app()["jd_text"]
    initial = docgen.normalize_cover_alignment_blueprint({
        "soft_skill_targets": [{"quote": "Collaborate across teams", "why": "explicit"}],
        "about_me_targets": [], "interest_targets": [], "technical_targets": [],
    }, jd)
    merged = docgen.merge_cover_alignment_blueprint(initial, {
        "soft_skill_targets": [
            {"quote": "Communicate clearly with stakeholders", "why": "omitted communication target"},
            {"quote": "Be a world-class leader", "why": "not in JD"},
        ],
        "about_me_targets": [], "interest_targets": [], "technical_targets": [],
    }, jd)
    assert [item["quote"] for item in merged["soft_skill_targets"]] == [
        "Collaborate across teams", "Communicate clearly with stakeholders"
    ]

VER = ROOT / "services" / "document-generation" / "verify_document_truth_v1.py"
VER_SPEC = importlib.util.spec_from_file_location("jobos_docverify_alignment_test", VER)
verifier = importlib.util.module_from_spec(VER_SPEC)
assert VER_SPEC and VER_SPEC.loader
VER_SPEC.loader.exec_module(verifier)


def test_truth_checker_allows_jd_bound_subjective_positioning_but_not_uncited_facts(monkeypatch):
    monkeypatch.setattr(verifier, "ollama_generate", lambda **kwargs: '{"verdict":"supported","reason":"subjective and JD-aligned"}')
    results = verifier.verify_claims(
        cur=None,
        claims=[{
            "claim": "I value collaborating across teams and keeping shared work clear.",
            "source_asset_id": None,
            "kind": "cover_letter_positioning",
            "positioning_kind": "soft_skill_positioning",
            "alignment_ids": ["S1"],
            "jd_requirement_quote": "Collaborate across teams",
        }],
        model="test", ollama_url="http://unused", timeout=1, temperature=0, num_ctx=512,
        verbose=False, valid_company_source_urls=set(), jd_text=_app()["jd_text"],
    )
    assert results[0]["verdict"] == "supported"

    results = verifier.verify_claims(
        cur=None,
        claims=[{
            "claim": "I administered Splunk in production for three years.",
            "source_asset_id": None,
            "kind": "cover_letter_evidence",
            "jd_requirement_quote": "Investigate security alerts",
        }],
        model="test", ollama_url="http://unused", timeout=1, temperature=0, num_ctx=512,
        verbose=False, valid_company_source_urls=set(), jd_text=_app()["jd_text"],
    )
    assert results[0]["verdict"] == "rule_violation"


def test_truth_checker_reaudits_experience_rewrite_against_asset_and_requirement_mapping(monkeypatch):
    asset = {
        "id": "asset-1", "title": "Internship evidence", "type": "source_document_asset",
        "summary": "During the internship, investigated alerts and documented findings for analyst review.",
        "bullets": "VERBATIM EMPLOYMENT SOURCE QUOTES (authoritative):\n- Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings for analyst review.",
        "positioning": "", "rules": [],
    }
    monkeypatch.setattr(verifier, "fetch_asset", lambda _cur, _asset_id: asset)
    monkeypatch.setattr(verifier, "ollama_generate", lambda **kwargs: '{"verdict":"supported","reason":"grounded and aligned","safe_rewrite":""}')
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 90,
        }
    }
    claim = {
        "claim": "Investigated security alerts and documented findings for analyst review.",
        "source_asset_id": "asset-1", "kind": "resume_experience_bullet_change", "slot": 1,
        "header_context": "Security Intern | Example Co | Jan-May 2025",
        "previous_bullet": "Investigated alerts and documented findings.",
        "jd_requirement_quote": "Investigate security alerts",
        "experience_evidence_quote": "Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings for analyst review",
        "matched_requirement_ids": ["M1"],
        "word_change_rationale": [
            {"before": "alerts", "after": "security alerts", "why": "Matches the JD wording without changing factual scope."},
            {"before": "documented findings", "after": "documented findings for analyst review", "why": "Uses the approved evidence's more specific review context."},
        ],
        "why_better": "Improves direct JD relevance while preserving the exact internship evidence boundary.",
    }
    results = verifier.verify_claims(
        cur=None, claims=[claim], model="test", ollama_url="http://unused", timeout=1,
        temperature=0, num_ctx=512, verbose=False, valid_company_source_urls=set(),
        jd_text=_app()["jd_text"], baseline_experience=baseline,
        matched_requirement_map={"M1": "Investigate security alerts"},
    )
    assert results[0]["verdict"] == "supported"

    bad = dict(claim, experience_evidence_quote="invented production ownership")
    results = verifier.verify_claims(
        cur=None, claims=[bad], model="test", ollama_url="http://unused", timeout=1,
        temperature=0, num_ctx=512, verbose=False, valid_company_source_urls=set(),
        jd_text=_app()["jd_text"], baseline_experience=baseline,
        matched_requirement_map={"M1": "Investigate security alerts"},
    )
    assert results[0]["verdict"] == "rule_violation"
    assert "absent from the cited approved user asset" in results[0]["reason"]


def test_final_cover_selected_fit_prompt_is_advisory_not_exhaustive():
    from services.common.document_prompt_templates_v1 import build_cover_letter_completeness_verifier_prompt
    app = _app()
    prompt = build_cover_letter_completeness_verifier_prompt(
        jd_text=app["jd_text"],
        cover_text="I value collaborating across teams.",
        alignment_blueprint={
            "about_me_targets": [],
            "interest_targets": [],
            "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
            "technical_targets": [],
        },
    )
    assert app["jd_text"] in prompt
    assert "FINISHED COVER LETTER" in prompt
    assert "Selected-Fit Auditor" in prompt
    assert "NOT an exhaustive A/I/S coverage gate" in prompt
    assert "do NOT need to mention every" in build_cover_letter_tailoring_prompt(
        app=app, asset_catalog="", alignment_blueprint={}
    )


def test_final_cover_selected_fit_audit_is_non_blocking(monkeypatch):
    blueprint = {
        "about_me_targets": [{"id": "A1", "quote": "Protect customer systems"}],
        "interest_targets": [],
        "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
        "technical_targets": [],
    }
    monkeypatch.setattr(verifier, "ollama_generate", lambda **kwargs: json.dumps({
        "verdict": "needs_polish",
        "reason": "the chosen positioning is generic",
    }))
    result = verifier.verify_cover_letter_completeness(
        jd_text=_app()["jd_text"], cover_text="I value collaboration.",
        alignment_blueprint=blueprint, model="test", ollama_url="http://unused",
        timeout=1, num_ctx=512,
    )
    assert result["verdict"] == "needs_polish"
    assert result["non_blocking"] is True

    monkeypatch.setattr(verifier, "ollama_generate", lambda **kwargs: "not-json")
    result = verifier.verify_cover_letter_completeness(
        jd_text=_app()["jd_text"], cover_text="I value collaboration.",
        alignment_blueprint=blueprint, model="test", ollama_url="http://unused",
        timeout=1, num_ctx=512,
    )
    assert result["verdict"] == "advisory_unavailable"
    assert result["non_blocking"] is True

def test_final_cover_selected_fit_audit_accepts_relevant_subset(monkeypatch):
    blueprint = {
        "about_me_targets": [{"id": "A1", "quote": "Protect customer systems"}],
        "interest_targets": [{"id": "I1", "quote": "Protect customer systems"}],
        "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
        "technical_targets": [],
    }
    monkeypatch.setattr(verifier, "ollama_generate", lambda **kwargs: json.dumps({
        "verdict": "supported",
        "missing_about_me_quotes": [],
        "missing_interest_quotes": [],
        "missing_soft_skill_quotes": [],
        "unmodeled_about_me_quotes": [],
        "unmodeled_interest_quotes": [],
        "unmodeled_soft_skill_quotes": [],
        "reason": "selected positioning is relevant and natural",
    }))
    result = verifier.verify_cover_letter_completeness(
        jd_text=_app()["jd_text"],
        cover_text=(
            "I am drawn to protecting customer systems. That mission is the part of the role that interests me most. "
            "I value collaborating across teams."
        ),
        alignment_blueprint=blueprint, model="test", ollama_url="http://unused",
        timeout=1, num_ctx=512,
    )
    assert result["verdict"] == "supported"


def test_truth_checker_accepts_general_jd_first_experience_rewrite_without_asset(monkeypatch):
    monkeypatch.setattr(
        verifier, "ollama_generate",
        lambda **kwargs: '{"verdict":"supported","reason":"role-plausible JD reframing","safe_rewrite":""}',
    )
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 100,
        }
    }
    claim = {
        "claim": "Investigated security alerts and documented findings for analyst review.",
        "source_asset_id": None,
        "kind": "resume_experience_bullet_change",
        "slot": 1,
        "header_context": baseline[1]["header_context"],
        "previous_bullet": baseline[1]["previous_bullet"],
        "jd_requirement_quote": "Investigate security alerts",
        "experience_evidence_quote": "",
        "matched_requirement_ids": [],
        "word_change_rationale": [
            {"before": "alerts", "after": "security alerts", "why": "Uses the target JD wording while staying inside the existing internship role."},
            {"before": "documented findings", "after": "documented findings for analyst review", "why": "Adds only a general role-plausible review framing, not a tool, metric, or outcome."},
        ],
        "why_better": "Improves target-JD relevance while preserving the immutable Security Intern identity and avoiding technical specifics.",
    }
    results = verifier.verify_claims(
        cur=None, claims=[claim], model="test", ollama_url="http://unused",
        timeout=1, temperature=0, num_ctx=512, verbose=False,
        valid_company_source_urls=set(), jd_text=_app()["jd_text"],
        baseline_experience=baseline, matched_requirement_map={"M1": "Investigate security alerts"},
        require_resume_alignment_ids=True,
    )
    assert results[0]["verdict"] == "supported"
    assert results[0]["source_asset_id"] is None
