from __future__ import annotations

import importlib.util
import os
import sys
import types
from datetime import date, datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

ROOT = Path(__file__).resolve().parents[1]


def load_path(name: str, relative: str, *, extra_path: Path | None = None):
    if extra_path is not None:
        sys.path.insert(0, str(extra_path))
    try:
        spec = importlib.util.spec_from_file_location(name, ROOT / relative)
        assert spec and spec.loader
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module
    finally:
        if extra_path is not None and sys.path and sys.path[0] == str(extra_path):
            sys.path.pop(0)


def test_docx_same_embedded_created_timestamp_but_changed_bytes_is_new_revision(tmp_path: Path):
    revisions = load_path("profile_source_revisions_test", "services/profile-ingestion/profile_source_revisions_v1.py")
    source_root = tmp_path / "data" / "profile_sources_v2"
    path = source_root / "00_official" / "resume.docx"
    path.parent.mkdir(parents=True)
    created = datetime(2026, 1, 2, 3, 4, 5)

    first = Document()
    first.core_properties.created = created
    first.core_properties.modified = created
    first.add_paragraph("Version one")
    first.save(path)
    meta1 = revisions.source_metadata(path, source_root)

    second = Document()
    second.core_properties.created = created
    second.core_properties.modified = created
    second.add_paragraph("Version two with changed facts")
    second.save(path)
    meta2 = revisions.source_metadata(path, source_root)

    assert meta1["embedded_created_at"] == meta2["embedded_created_at"]
    assert meta1["content_sha256"] != meta2["content_sha256"]
    assert meta1["logical_source_key"] == meta2["logical_source_key"]


def test_malformed_pdf_creation_and_mod_dates_are_non_authoritative_and_do_not_crash(tmp_path: Path):
    revisions = load_path("profile_source_revisions_pdf_test", "services/profile-ingestion/profile_source_revisions_v1.py")
    path = tmp_path / "broken_metadata.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_metadata({"/CreationDate": "D:not-a-date", "/ModDate": "totally malformed", "/Title": "Fixture"})
    with path.open("wb") as stream:
        writer.write(stream)

    metadata = revisions.embedded_metadata(path)
    assert metadata["embedded_created_at"] is None
    assert metadata["embedded_modified_at"] is None
    assert metadata["embedded_title"] == "Fixture"
    assert metadata["pdf_page_count"] == 1


def _verified(value, display: str, when: date, *, show=True, status="user_verified", expires_at=None):
    return {
        "value": value,
        "display_value": display,
        "verification_status": status,
        "show_on_resume": show,
        "verified_at": datetime(when.year, when.month, when.day, tzinfo=timezone.utc),
        "expires_at": expires_at,
    }


def complete_fixed_fields(policy, *, today: date) -> dict[str, dict]:
    fields: dict[str, dict] = {}
    for definition in policy.FIELD_DEFINITIONS:
        if definition.value_type == "bool":
            value, display = False, "No"
        elif definition.key == "education.graduation_date":
            value, display = "2026-09", "2026-09"
        else:
            value, display = "fixture", "fixture"
        fields[definition.key] = _verified(value, display, today, show=definition.show_on_resume_default)
    fields["education.gpa.show_on_resume"] = _verified(False, "No", today, show=False)
    fields["certifications.reviewed"] = _verified(True, "Yes", today, show=False)
    return fields


def test_fixed_fields_conflicts_gpa_staleness_and_certification_lifecycle():
    from services.common import fixed_profile_policy as policy

    today = date(2026, 8, 24)
    fields = complete_fixed_fields(policy, today=today)
    ready = policy.readiness_from_records(fields, [], today=today)
    assert ready["fixed_fields_ready"] is True

    fields["personal.email"] = _verified("old@example.com", "old@example.com", date(2025, 1, 1))
    stale = policy.readiness_from_records(fields, [], today=today)
    assert "personal.email" in stale["stale_fields"]
    assert stale["fixed_fields_ready"] is False

    fields = complete_fixed_fields(policy, today=today)
    fields["education.gpa.show_on_resume"] = _verified(True, "Yes", today, show=False)
    fields["education.gpa.value"] = _verified("3.8", "3.8", date(2026, 1, 1))
    fields["education.gpa.scale"] = _verified("4.0", "4.0", date(2026, 1, 1))
    fields["education.gpa.status"] = _verified("current", "current", date(2026, 1, 1))
    gpa_stale = policy.readiness_from_records(fields, [], today=today)
    assert {"education.gpa.value", "education.gpa.scale", "education.gpa.status"}.issubset(gpa_stale["stale_fields"])

    fields = complete_fixed_fields(policy, today=today)
    bad_cert = [{
        "name": "Fixture Cert", "show_on_resume": True,
        "certification_status": "expired", "verification_status": "user_verified",
        "expires_at": date(2026, 1, 1),
    }]
    cert_report = policy.readiness_from_records(fields, bad_cert, today=today)
    assert cert_report["invalid_visible_certifications"] == ["Fixture Cert"]
    assert cert_report["fixed_fields_ready"] is False

    fields["education.degree"]["verification_status"] = "conflict"
    conflict = policy.readiness_from_records(fields, [], today=today)
    assert "education.degree" in conflict["conflicting_fields"]
    assert conflict["fixed_fields_ready"] is False


def test_fixed_fields_reject_gpa_above_scale_and_invalid_cert_dates():
    from services.common import fixed_profile_policy as policy

    today = date(2026, 8, 24)
    fields = complete_fixed_fields(policy, today=today)
    fields["education.gpa.show_on_resume"] = _verified(True, "Yes", today, show=False)
    fields["education.gpa.value"] = _verified("5.0", "5.0", today)
    fields["education.gpa.scale"] = _verified("4.0", "4.0", today)
    fields["education.gpa.status"] = _verified("final", "final", today)
    report = policy.readiness_from_records(fields, [], today=today)
    assert report["fixed_fields_ready"] is False
    assert "education.gpa.value" in report["invalid_fields"]

    fields["education.gpa.value"] = _verified("3.8", "3.8", today)
    future_cert = [{
        "name": "Future Cert", "show_on_resume": True,
        "certification_status": "earned", "verification_status": "user_verified",
        "earned_at": date(2026, 9, 1), "expires_at": date(2027, 9, 1),
    }]
    cert_report = policy.readiness_from_records(fields, future_cert, today=today)
    assert cert_report["fixed_fields_ready"] is False
    assert cert_report["invalid_visible_certifications"] == ["Future Cert"]


def test_fixed_field_document_extractor_only_accepts_explicit_labels():
    fixed = load_path("fixed_profile_fields_test", "services/profile-ingestion/fixed_profile_fields_v1.py")
    payload = fixed.candidate_suggestions_from_text(
        """
        GPA: 3.75 / 4.0
        Degree: MSc International Business
        Certification: AWS Cloud Practitioner
        I hope to earn another certification someday.
        """
    )
    by_key = {item["field_key"]: item["value"] for item in payload["fields"]}
    assert by_key["education.gpa.value"] == "3.75"
    assert by_key["education.gpa.scale"] == "4.0"
    assert by_key["education.degree"] == "MSc International Business"
    assert payload["certifications"] == ["AWS Cloud Practitioner"]


def load_repo_freshness():
    repo_audit = ROOT / "services" / "repo-audit"
    return load_path("repository_freshness_test", "services/repo-audit/repository_freshness_v1.py", extra_path=repo_audit)


def load_repo_claims():
    return load_path("repository_claims_test", "services/repo-audit/repository_claims_v1.py")


def test_github_same_head_is_zero_diff_without_network(monkeypatch):
    fresh = load_repo_freshness()
    called = False

    def forbidden(*_args, **_kwargs):
        nonlocal called
        called = True
        raise AssertionError("network must not be needed for identical SHAs")

    monkeypatch.setattr(fresh, "github_json", forbidden)
    sha = "a" * 40
    files, full, reason = fresh.github_change_set("owner/repo", sha, sha, None)
    assert files == []
    assert full is False
    assert reason is None
    assert called is False


def test_github_readme_only_vs_implementation_diff_classification():
    claims = load_repo_claims()
    docs = claims.classify_changed_files([{"filename": "README.md", "status": "modified"}])
    runtime = claims.classify_changed_files([{"filename": "services/api.py", "status": "modified"}])
    generated = claims.classify_changed_files([{"filename": "dist/app.min.js", "status": "modified"}])
    assert docs["requires_analysis"] is False
    assert docs["buckets"]["documentation"] == ["README.md"]
    assert runtime["requires_analysis"] is True
    assert runtime["buckets"]["runtime"] == ["services/api.py"]
    assert generated["requires_analysis"] is False


def test_github_analysis_plan_skips_readme_only_but_analyzes_runtime_and_first_snapshot():
    fresh = load_repo_freshness()
    claims = load_repo_claims()
    docs = claims.classify_changed_files([{"filename": "README.md", "status": "modified"}])
    runtime = claims.classify_changed_files([{"filename": "services/api.py", "status": "modified"}])
    assert fresh._requires_material_analysis(prior_sha="a" * 40, full_reanalysis=False, classification=docs) is False
    assert fresh._requires_material_analysis(prior_sha="a" * 40, full_reanalysis=False, classification=runtime) is True
    assert fresh._requires_material_analysis(prior_sha=None, full_reanalysis=True, classification=docs) is True


def test_github_history_rewrite_compare_failure_forces_full_reanalysis(monkeypatch):
    fresh = load_repo_freshness()

    def broken_compare(*_args, **_kwargs):
        raise fresh.RefreshError("comparison not available")

    monkeypatch.setattr(fresh, "github_compare", broken_compare)
    files, full, reason = fresh.github_change_set("owner/repo", "a" * 40, "b" * 40, None)
    assert files == []
    assert full is True
    assert reason.startswith("compare_unavailable:")


def test_project_asset_versioning_reuses_same_material_and_versions_only_changed_material():
    fresh = load_repo_freshness()
    h1 = "a" * 64
    h2 = "b" * 64
    assert fresh._asset_version_decision("approved", h1, h1) == "revalidate_approved"
    assert fresh._asset_version_decision("needs_review", h1, h1) == "keep_candidate"
    assert fresh._asset_version_decision("rejected", h1, h1) == "keep_terminal"
    assert fresh._asset_version_decision("approved", h1, h2) == "create_candidate"
    assert fresh._asset_version_decision(None, None, h1) == "create_candidate"


def test_project_source_material_hash_is_snapshot_independent_but_material_sensitive():
    fresh = load_repo_freshness()
    material = {
        "title": "ApplyOps",
        "canonical": "Current implementation",
        "summary": "summary",
        "resume": "- claim",
        "cover": "purpose",
        "tools": ["Python"],
        "rules": ["no overclaim"],
        "authority": {"implementation": {"github": 0.7, "document": 0.3}},
    }
    first = fresh._material_hash(material)
    # HEAD is intentionally not an argument: a README-only commit may advance
    # source_snapshot_hash without manufacturing a new material version.
    assert first == fresh._material_hash(dict(material))
    changed = dict(material)
    changed["tools"] = ["Python", "Fastapi"]
    assert first != fresh._material_hash(changed)


def test_repository_claims_ignore_python_comments_and_docstrings(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / "app.py").write_text(
        "\"\"\"Planned: fail closed, approval review, reconciliation and SHA-256.\"\"\"\n"
        "# TODO: implement fail closed approval review reconciliation and SHA-256\n"
        "print('hello')  # approval review should be added later\n",
        encoding="utf-8",
    )
    extracted = claims.extract_claims(repo)
    assert not [item for item in extracted if item["claim_key"].startswith("control:")]


def test_repository_claims_still_detect_real_python_controls(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / "worker.py").write_text(
        "import hashlib\n"
        "digest = hashlib.sha256(b'x').hexdigest()\n"
        "needs_reconciliation = True\n"
        "approval = True\n"
        "sql = 'SELECT id FROM tasks FOR UPDATE'\n",
        encoding="utf-8",
    )
    keys = {item["claim_key"] for item in claims.extract_claims(repo)}
    assert "control:sha256_integrity" in keys
    assert "control:reconciliation" in keys
    assert "control:approval_review" in keys
    assert "control:row_locking" in keys


def test_repository_test_surface_requires_test_filename_or_directory(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / "contest.py").write_text("print('not a test')\n", encoding="utf-8")
    assert "surface:automated_tests" not in {item["claim_key"] for item in claims.extract_claims(repo)}
    (repo / "tests").mkdir()
    (repo / "tests" / "test_worker.py").write_text("def test_ok(): assert True\n", encoding="utf-8")
    assert "surface:automated_tests" in {item["claim_key"] for item in claims.extract_claims(repo)}


def test_last_known_good_is_rejected_after_newer_unanalyzed_snapshot_is_observed():
    fresh = load_repo_freshness()
    assert fresh._last_known_good_eligible(
        has_last_analyzed=True, age_hours=1.0, snapshot_is_current=True, max_stale_hours=24
    ) is True
    assert fresh._last_known_good_eligible(
        has_last_analyzed=True, age_hours=1.0, snapshot_is_current=False, max_stale_hours=24
    ) is False
    assert fresh._last_known_good_eligible(
        has_last_analyzed=True, age_hours=25.0, snapshot_is_current=True, max_stale_hours=24
    ) is False


def test_offline_repository_claims_are_file_line_and_blob_pinned(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    (repo / "services").mkdir(parents=True)
    source = repo / "services" / "api.py"
    source.write_text("from fastapi import FastAPI\napp = FastAPI()\n", encoding="utf-8")
    extracted = claims.extract_claims(repo)
    tech = next(item for item in extracted if item["claim_key"] == "tech:fastapi")
    assert tech["evidence_path"] == "services/api.py"
    assert tech["source_line_start"] >= 1
    assert tech["source_line_end"] >= tech["source_line_start"]
    assert len(tech["evidence_blob_sha"]) == 64  # exact file SHA-256


def load_generator(monkeypatch):
    os.environ.setdefault("JOBOS_DB_PASSWORD", "test-only")
    psycopg = types.ModuleType("psycopg")
    psycopg.connect = lambda *_a, **_k: None
    psycopg.Error = Exception
    psycopg_types = types.ModuleType("psycopg.types")
    psycopg_json = types.ModuleType("psycopg.types.json")
    psycopg_json.Jsonb = lambda value: value
    monkeypatch.setitem(sys.modules, "psycopg", psycopg)
    monkeypatch.setitem(sys.modules, "psycopg.types", psycopg_types)
    monkeypatch.setitem(sys.modules, "psycopg.types.json", psycopg_json)
    return load_path("generate_documents_freshness_test", "services/document-generation/generate_documents_v1.py")


def test_resume_generator_live_freshness_preflight_blocks_before_database(monkeypatch, capsys):
    generator = load_generator(monkeypatch)
    monkeypatch.setattr(generator, "run_live_project_freshness", lambda **_kw: (False, "GitHub project stale"))
    monkeypatch.setattr(sys, "argv", ["generate_documents_v1.py", "--application-id", "fixture", "--doc-type", "resume"])
    assert generator.main() == 2
    output = capsys.readouterr().out
    assert "live project freshness preflight failed" in output
    assert "GitHub project stale" in output


def test_resume_database_freshness_wrapper_returns_blockers(monkeypatch):
    generator = load_generator(monkeypatch)
    import services.common.profile_freshness as profile_freshness

    monkeypatch.setattr(profile_freshness, "assess_resume_profile", lambda _cur, **_kw: {"resume_profile_ready": False})
    monkeypatch.setattr(profile_freshness, "explain_blockers", lambda _report: ["fixture blocker"])
    ready, report, blockers = generator.database_resume_freshness(object())
    assert ready is False
    assert report == {"resume_profile_ready": False}
    assert blockers == ["fixture blocker"]

class _FreshnessCursor:
    def __init__(self, *, approved_project_asset: bool):
        self._rows = []
        self._one = None
        self.approved_project_asset = approved_project_asset

    def execute(self, sql, params=None):
        q = " ".join(str(sql).split()).casefold()
        # Fixed fields: return a complete fresh set.
        if "from candidate_fixed_fields order by field_key" in q:
            from services.common import fixed_profile_policy as policy
            today = date(2026, 8, 24)
            fields = complete_fixed_fields(policy, today=today)
            self._rows = [
                (key, row["value"], row["display_value"], row["verification_status"],
                 row["show_on_resume"], "candidate", row["verified_at"], row["expires_at"])
                for key, row in fields.items()
            ]
            self._one = None
        elif "from candidate_certifications order by" in q:
            self._rows = []
            self._one = None
        elif "candidate_fixed_field_suggestions" in q:
            self._one = (0,)
            self._rows = []
        elif "from repository_evidence_sources" in q and "repo_full_name = any" in q:
            self._rows = [("owner/repo", "fresh", "confirmed_by_user", "snap-1", "snap-1", "a" * 40, None)]
            self._one = None
        elif "select count(*) from repository_claims" in q:
            self._one = (0,)
            self._rows = []
        elif "select count(*) from project_source_conflicts" in q:
            self._one = (0,)
            self._rows = []
        elif "select distinct project_id from profile_assets" in q:
            self._rows = [("applyops",)] if self.approved_project_asset else []
            self._one = None
        elif "freshness_status not in ('fresh','not_applicable')" in q:
            self._one = (0,)
            self._rows = []
        elif "status in ('needs_review','pending_review','draft')" in q:
            self._one = (0,)
            self._rows = []
        elif "select count(*) from profile_briefs" in q:
            self._one = (1,)
            self._rows = []
        elif "select count(*) from profile_context_packs" in q:
            self._one = (1,)
            self._rows = []
        else:
            raise AssertionError(f"Unhandled freshness SQL: {sql}")

    def fetchall(self):
        return list(self._rows)

    def fetchone(self):
        if self._one is not None:
            return self._one
        if self._rows:
            return self._rows[0]
        raise AssertionError("fetchone called with no prepared result")


def test_resume_profile_requires_one_approved_current_authority_asset_per_configured_project(monkeypatch):
    import services.common.profile_freshness as freshness
    monkeypatch.setattr(freshness, "load_registry", lambda: {
        "projects": [{
            "project_id": "applyops", "dynamic_source_mode": "github_primary",
            "github_repo_full_name": "owner/repo",
        }]
    })
    blocked = freshness.assess_resume_profile(_FreshnessCursor(approved_project_asset=False), today=date(2026, 8, 24))
    assert blocked["resume_profile_ready"] is False
    assert blocked["missing_approved_project_assets"] == ["applyops"]
    assert any("approved current authority asset" in item for item in freshness.explain_blockers(blocked))

    ready = freshness.assess_resume_profile(_FreshnessCursor(approved_project_asset=True), today=date(2026, 8, 24))
    assert ready["missing_approved_project_assets"] == []
    assert ready["resume_profile_ready"] is True


def test_readme_prose_cannot_prove_security_implementation(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / "README.md").write_text(
        "The system has human approval, fail-closed reconciliation, and SHA-256 integrity.\n",
        encoding="utf-8",
    )
    extracted = claims.extract_claims(repo)
    assert not [item for item in extracted if item["claim_kind"] in {"security_control", "reliability_control", "concurrency_control"}]


def test_security_claim_requires_implementation_source(tmp_path: Path):
    claims = load_repo_claims()
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / "worker.py").write_text("approval = True\nraise PermanentTaskError('blocked')\n", encoding="utf-8")
    extracted = claims.extract_claims(repo)
    keys = {item["claim_key"] for item in extracted}
    assert "control:approval_review" in keys
    assert "control:fail_closed" in keys


def test_cached_snapshot_rejects_dirty_worktree(tmp_path: Path):
    import subprocess
    fresh = load_repo_freshness()
    origin = tmp_path / "origin"
    origin.mkdir()
    subprocess.run(["git", "init", "-q"], cwd=origin, check=True)
    subprocess.run(["git", "config", "user.email", "fixture@example.com"], cwd=origin, check=True)
    subprocess.run(["git", "config", "user.name", "Fixture"], cwd=origin, check=True)
    (origin / "app.py").write_text("print(1)\n", encoding="utf-8")
    subprocess.run(["git", "add", "."], cwd=origin, check=True)
    subprocess.run(["git", "commit", "-qm", "fixture"], cwd=origin, check=True)
    head = subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=origin, text=True).strip()
    fresh.SNAPSHOT_ROOT = tmp_path / "snapshots"
    checkout = fresh.ensure_immutable_checkout(repo_full_name="owner/repo", clone_url=str(origin), head_sha=head, token=None)
    (checkout / "app.py").write_text("print('tampered')\n", encoding="utf-8")
    refreshed = fresh.ensure_immutable_checkout(repo_full_name="owner/repo", clone_url=str(origin), head_sha=head, token=None)
    assert (refreshed / "app.py").read_text(encoding="utf-8") == "print(1)\n"


class _LastKnownGoodCursor(_FreshnessCursor):
    def execute(self, sql, params=None):
        q = " ".join(str(sql).split()).casefold()
        if "extract(epoch from (now() - snap.analyzed_at))" in q:
            self._rows = [("owner/repo", 2.0)]
            self._one = None
        elif "from repository_evidence_sources" in q and "repo_full_name = any" in q:
            self._rows = [("owner/repo", "unavailable", "confirmed_by_user", "snap-1", "snap-1", "a" * 40, "temporary network failure")]
            self._one = None
        else:
            super().execute(sql, params)


def test_database_gate_accepts_bounded_last_known_good_only_when_live_gate_allows_it(monkeypatch):
    import services.common.profile_freshness as freshness
    monkeypatch.setattr(freshness, "load_registry", lambda: {
        "projects": [{"project_id": "applyops", "dynamic_source_mode": "github_primary", "github_repo_full_name": "owner/repo"}]
    })
    strict = freshness.assess_resume_profile(_LastKnownGoodCursor(approved_project_asset=True), today=date(2026, 8, 24))
    assert strict["resume_profile_ready"] is False
    assert strict["stale_repository_sources"] == ["owner/repo"]
    bounded = freshness.assess_resume_profile(
        _LastKnownGoodCursor(approved_project_asset=True), today=date(2026, 8, 24), allow_last_known_good_hours=24
    )
    assert bounded["stale_repository_sources"] == []
    assert bounded["last_known_good_repository_sources"] == ["owner/repo"]
    assert bounded["resume_profile_ready"] is True


class _DocumentOnlyCursor(_FreshnessCursor):
    def execute(self, sql, params=None):
        q = " ".join(str(sql).split()).casefold()
        if "source_strategy='project_document_only_v1'" in q and "select distinct project_id" in q:
            self._rows = [("applyops",)] if self.approved_project_asset else []
            self._one = None
        elif "project_id = any" in q and "freshness_status not in ('fresh','not_applicable')" in q:
            self._one = (0,)
            self._rows = []
        elif "project_id = any" in q and "status in ('needs_review','pending_review','draft')" in q:
            self._one = (0,)
            self._rows = []
        else:
            super().execute(sql, params)


def test_document_only_project_requires_approved_document_asset(monkeypatch):
    import services.common.profile_freshness as freshness
    monkeypatch.setattr(freshness, "load_registry", lambda: {
        "projects": [{"project_id": "applyops", "dynamic_source_mode": "document_only", "github_repo_full_name": ""}]
    })
    blocked = freshness.assess_resume_profile(_DocumentOnlyCursor(approved_project_asset=False), today=date(2026, 8, 24))
    assert blocked["missing_document_only_project_assets"] == ["applyops"]
    assert blocked["resume_profile_ready"] is False
    ready = freshness.assess_resume_profile(_DocumentOnlyCursor(approved_project_asset=True), today=date(2026, 8, 24))
    assert ready["missing_document_only_project_assets"] == []
    assert ready["resume_profile_ready"] is True
