"""Acceptance-level behavior checks for JD-targeted documents.

These tests exercise the policy boundary across generation validation, truth QA,
and fixed-template rendering without requiring PostgreSQL/Ollama.
"""
from __future__ import annotations

import importlib.util
import os
from pathlib import Path
import sys
import types

from docx import Document

os.environ.setdefault("JOBOS_DB_PASSWORD", "unit-test")

if "psycopg" not in sys.modules:
    psycopg = types.ModuleType("psycopg")
    psycopg.connect = lambda *_a, **_k: None
    psycopg_types = types.ModuleType("psycopg.types")
    psycopg_json = types.ModuleType("psycopg.types.json")
    psycopg_json.Jsonb = lambda value: value
    sys.modules["psycopg"] = psycopg
    sys.modules["psycopg.types"] = psycopg_types
    sys.modules["psycopg.types.json"] = psycopg_json

ROOT = Path(__file__).resolve().parents[1]


def _load(name: str, relative: str):
    spec = importlib.util.spec_from_file_location(name, ROOT / relative)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


docgen = _load("jobos_tailoring_acceptance_docgen", "services/document-generation/generate_documents_v1.py")
verifier = _load("jobos_tailoring_acceptance_verify", "services/document-generation/verify_document_truth_v1.py")
renderer = _load("jobos_tailoring_acceptance_renderer", "services/document-generation/resume_template_renderer.py")

JD = (
    "Investigate security alerts. Automate reporting with Python. "
    "Communicate clearly with stakeholders. Collaborate across teams. "
    "Learn new security tooling quickly. Protect customer systems."
)


def _experience_change() -> dict:
    return {
        "slot": 1,
        "header_context": "Security Intern | Example Co | Jan-May 2025",
        "previous_bullet": "Investigated alerts and documented findings.",
        "text": "Investigated security alerts and documented findings for analyst review.",
        "source_asset_id": "resume-asset",
        "jd_requirement_quote": "Investigate security alerts",
        "experience_evidence_quote": (
            "Security Intern | Example Co | Jan-May 2025 investigated alerts and documented findings for analyst review"
        ),
        "matched_requirement_ids": ["M1"],
        "word_change_rationale": [
            {
                "before": "alerts",
                "after": "security alerts",
                "why": "Uses the JD's precise security-alert wording without changing the evidenced scope.",
            },
            {
                "before": "documented findings",
                "after": "documented findings for analyst review",
                "why": "Uses the approved employment evidence's more specific review context.",
            },
        ],
        "why_better": "Directly matches the alert-investigation requirement while preserving the internship evidence boundary.",
    }


def _skill(item: str, requirement_id: str, jd_quote: str) -> dict:
    return {
        "category": "Relevant Skills",
        "items": item,
        "source_asset_id": "profile-asset",
        "matched_requirement_ids": [requirement_id],
        "jd_requirement_quote": jd_quote,
        "skill_evidence_quote": f"approved evidence for {item}",
    }


def _make_resume_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("CANDIDATE NAME | candidate@example.com")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Security Intern | Example Co | Jan-May 2025")
    doc.add_paragraph("Investigated alerts and documented findings.", style="List Paragraph")
    doc.add_paragraph("PROJECTS")
    for i in range(6):
        doc.add_paragraph(f"Project {i + 1} — baseline subtitle {i + 1} | GitHub\tJan 2026")
        doc.add_paragraph(f"baseline project bullet {2*i+1}", style="List Paragraph")
        doc.add_paragraph(f"baseline project bullet {2*i+2}", style="List Paragraph")
    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("SKILLS")
    for category in ("Network", "Security", "AI", "Vision", "Programming"):
        p = doc.add_paragraph()
        p.add_run(f"{category}: ").bold = True
        p.add_run("baseline")
    doc.save(path)


def test_resume_acceptance_hits_80_percent_full_jd_and_changes_only_experience_bullet(tmp_path):
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 100,
        }
    }
    parsed = {
        "experience_updates": [_experience_change()],
        "skill_lines_ranked": [
            _skill("Python reporting automation", "M2", "Automate reporting with Python"),
            _skill("Stakeholder communication", "M3", "Communicate clearly with stakeholders"),
            _skill("Cross-team collaboration", "M4", "Collaborate across teams"),
            _skill("Rapid security-tool learning", "M5", "Learn new security tooling quickly"),
        ],
        "not_supported": ["Protect customer systems"],
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "resume",
        parsed,
        {"resume-asset", "profile-asset"},
        jd_text=JD,
        experience_baselines=baseline,
        experience_source_asset_ids={"resume-asset"},
        matched_requirement_ids={"M1", "M2", "M3", "M4", "M5"},
        resume_coverage_target_percent=80,
        resume_total_material_requirement_count=6,
    )
    assert content
    assert evidence["jd_alignment"]["coverage_percent"] == 83.3
    assert evidence["jd_alignment"]["gate_passed"] is True
    assert set(used) == {"resume-asset", "profile-asset"}
    assert evidence["jd_alignment"]["target_met"] is True
    assert not any("below the 80% target" in item for item in evidence["warnings"])

    template = tmp_path / "template.docx"
    output = tmp_path / "resume.docx"
    _make_resume_template(template)
    renderer.render_docx(
        template=template,
        output=output,
        experience_bullets=evidence["resume_template"]["experience_bullets"],
        project_bullets=[],
        skill_lines=[],
        project_subtitles=[],
    )
    rendered = [p.text for p in Document(output).paragraphs]
    assert "Security Intern | Example Co | Jan-May 2025" in rendered
    assert "Investigated security alerts and documented findings for analyst review." in rendered
    assert "Investigated alerts and documented findings." not in rendered


def test_resume_acceptance_uses_truthful_ceiling_instead_of_fabricating_80_percent():
    parsed = {
        "skill_lines_ranked": [
            _skill("Alert triage", "M1", "Investigate security alerts"),
            _skill("Python reporting automation", "M2", "Automate reporting with Python"),
            _skill("Stakeholder communication", "M3", "Communicate clearly with stakeholders"),
        ]
    }
    content, _used, evidence, _dropped = docgen.validate_and_render(
        "resume", parsed, {"profile-asset"}, jd_text=JD,
        matched_requirement_ids={"M1", "M2", "M3"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=5,
    )
    assert content
    assert evidence["jd_alignment"]["truthful_coverage_ceiling_percent"] == 60.0
    assert evidence["jd_alignment"]["target_reachable_truthfully"] is False
    assert evidence["jd_alignment"]["covered_requirement_count"] == 3
    assert evidence["jd_alignment"]["gate_passed"] is True


def test_cover_letter_acceptance_selects_best_ais_and_rejects_ungrounded_technical(monkeypatch):
    blueprint = {
        "about_me_targets": [{"id": "A1", "quote": "Protect customer systems"}],
        "interest_targets": [{"id": "I1", "quote": "Protect customer systems"}],
        "soft_skill_targets": [{"id": "S1", "quote": "Collaborate across teams"}],
        "technical_targets": [{"id": "T1", "quote": "Automate reporting with Python"}],
    }
    parsed = {
        "paragraphs": [
            {"sentences": [
                {"kind": "about_me_positioning", "text": "I am drawn to work centered on protecting customer systems.", "alignment_ids": ["A1"], "jd_requirement_quote": "Protect customer systems", "source_asset_id": "none"},
                {"kind": "role_interest", "text": "Protecting customer systems is the part of this role that interests me most.", "alignment_ids": ["I1"], "jd_requirement_quote": "Protect customer systems", "source_asset_id": "none"},
                {"kind": "soft_skill_positioning", "text": "I value collaborating across teams and keeping shared work clear.", "alignment_ids": ["S1"], "jd_requirement_quote": "Collaborate across teams", "source_asset_id": "none"},
            ]},
            {"sentences": [
                {"kind": "technical_evidence", "text": "I used Python to automate recurring reporting.", "alignment_ids": ["T1"], "jd_requirement_quote": "Automate reporting with Python", "source_asset_id": "asset-1", "candidate_evidence_quote": "automated recurring reporting with Python"},
            ]},
        ]
    }
    content, _used, evidence, _dropped = docgen.validate_and_render(
        "cover_letter", parsed, {"asset-1"}, jd_text=JD,
        company="Example Co", job_title="Security Analyst", cover_alignment_blueprint=blueprint,
    )
    assert content
    assert evidence["cover_alignment"]["coverage_percent"] == 100.0

    asset = {
        "id": "asset-1", "title": "Approved automation evidence", "type": "project_asset",
        "summary": "automated recurring reporting with Python", "bullets": "", "positioning": "", "rules": [],
    }
    monkeypatch.setattr(verifier, "fetch_asset", lambda _cur, _asset_id: asset)
    monkeypatch.setattr(verifier, "ollama_generate", lambda **_kwargs: '{"verdict":"supported","reason":"grounded","safe_rewrite":""}')
    technical = next(c for c in evidence["claims"] if c.get("kind") == "cover_letter_evidence")
    results = verifier.verify_claims(
        cur=None, claims=[technical], model="test", ollama_url="http://unused",
        timeout=1, temperature=0, num_ctx=512, verbose=False,
        valid_company_source_urls=set(), jd_text=JD,
    )
    assert results[0]["verdict"] == "supported"

    bad = dict(technical, candidate_evidence_quote="administered Splunk in production for three years")
    results = verifier.verify_claims(
        cur=None, claims=[bad], model="test", ollama_url="http://unused",
        timeout=1, temperature=0, num_ctx=512, verbose=False,
        valid_company_source_urls=set(), jd_text=JD,
    )
    assert results[0]["verdict"] == "rule_violation"

    parsed["paragraphs"][0]["sentences"] = parsed["paragraphs"][0]["sentences"][:2]
    content, _used, evidence, dropped = docgen.validate_and_render(
        "cover_letter", parsed, {"asset-1"}, jd_text=JD,
        company="Example Co", job_title="Security Analyst", cover_alignment_blueprint=blueprint,
    )
    assert content
    assert evidence["cover_alignment"]["coverage_percent"] == 100.0
    assert "S1" in evidence["cover_alignment"]["available_positioning_ids"]
    assert "S1" not in evidence["cover_alignment"]["selected_positioning_ids"]
    assert not any("missing targets" in item for item in dropped)


def test_sparse_information_soft_degrades_instead_of_hard_failing():
    resume_content, resume_used, resume_evidence, resume_dropped = docgen.validate_and_render(
        "resume", {}, set(), jd_text=JD,
        experience_baselines={
            1: {
                "slot": 1,
                "header_context": "Security Intern | Example Co | Jan-May 2025",
                "previous_bullet": "Investigated alerts and documented findings.",
                "max_chars": 100,
            }
        },
        matched_requirement_ids={"M1", "M2", "M3"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=6,
    )
    assert "Resume template preserved" in resume_content
    assert resume_used == []
    assert resume_evidence["jd_alignment"]["gate_passed"] is True
    assert resume_evidence["jd_alignment"]["target_met"] is False
    assert resume_evidence["warnings"]

    cover_content, cover_used, cover_evidence, cover_dropped = docgen.validate_and_render(
        "cover_letter", {"paragraphs": []}, set(), jd_text=JD,
        company="Example Co", job_title="Security Analyst", cover_alignment_blueprint={},
    )
    assert "I am applying for the Security Analyst position at Example Co." in cover_content
    assert "welcome the opportunity to contribute" in cover_content
    assert cover_used == []
    assert cover_evidence["warnings"]
    assert cover_dropped == []


def test_experience_rewrite_can_be_jd_first_without_official_resume_source():
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 100,
        }
    }
    parsed = {
        "experience_updates": [{
            "slot": 1,
            "header_context": baseline[1]["header_context"],
            "previous_bullet": baseline[1]["previous_bullet"],
            "text": "Investigated security alerts and documented findings for analyst review.",
            "source_asset_id": "none",
            "jd_requirement_quote": "Investigate security alerts",
            "experience_evidence_quote": "",
            "matched_requirement_ids": ["M1"],
            "word_change_rationale": [
                {"before": "alerts", "after": "security alerts", "why": "Matches the target JD while preserving the existing analyst internship scope."},
                {"before": "documented findings", "after": "documented findings for analyst review", "why": "Keeps a general review framing consistent with the immutable title."},
            ],
            "why_better": "Prioritizes the target JD and remains compatible with the immutable Security Intern job title.",
        }]
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "resume", parsed, set(), jd_text=JD, experience_baselines=baseline,
        matched_requirement_ids={"M1"}, resume_coverage_target_percent=80,
        resume_total_material_requirement_count=6,
    )
    assert "Investigated security alerts" in content
    assert used == []
    assert evidence["resume_template"]["experience_bullets"][0]["source_asset_id"] is None
    assert not dropped


def test_experience_jd_rewrite_without_supported_requirement_id_is_kept_but_not_counted():
    baseline = {
        1: {
            "slot": 1,
            "header_context": "Security Intern | Example Co | Jan-May 2025",
            "previous_bullet": "Investigated alerts and documented findings.",
            "max_chars": 100,
        }
    }
    parsed = {
        "experience_updates": [{
            "slot": 1,
            "header_context": baseline[1]["header_context"],
            "previous_bullet": baseline[1]["previous_bullet"],
            "text": "Supported security review workflows by investigating alerts and documenting findings.",
            "source_asset_id": "none",
            "jd_requirement_quote": "Protect customer systems",
            "experience_evidence_quote": "",
            "matched_requirement_ids": [],
            "word_change_rationale": [
                {"before": "Investigated alerts", "after": "Supported security review workflows by investigating alerts", "why": "Connects the existing internship work to the JD's protection mission without adding a technical tool or outcome."},
                {"before": "documented findings", "after": "documenting findings", "why": "Keeps the original documentation responsibility while improving sentence flow."},
            ],
            "why_better": "Makes the existing Security Intern responsibility more relevant to the target JD without claiming unsupported technical specifics.",
        }]
    }
    content, used, evidence, dropped = docgen.validate_and_render(
        "resume", parsed, set(), jd_text=JD, experience_baselines=baseline,
        matched_requirement_ids={"M1", "M2", "M3"},
        resume_coverage_target_percent=80, resume_total_material_requirement_count=6,
    )
    assert "Supported security review workflows" in content
    assert used == []
    assert evidence["jd_alignment"]["covered_requirement_count"] == 0
    assert evidence["jd_alignment"]["target_met"] is False
    assert any("without counting it toward resume coverage" in w for w in evidence["warnings"])
    assert dropped == []
