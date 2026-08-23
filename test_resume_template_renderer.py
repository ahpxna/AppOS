"""Template slot tests; PDF conversion is verified separately where Office exists."""

import importlib.util
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
PATH = ROOT / "services" / "document-generation" / "resume_template_renderer.py"
SPEC = importlib.util.spec_from_file_location("jobos_resume_template_renderer_test", PATH)
renderer = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(renderer)


def make_template(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("CANDIDATE NAME | github.com/ahpxna")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("PROJECTS")
    for number in range(6):
        doc.add_paragraph(f"Project {number + 1} — original subtitle {number + 1} | GitHub\tJan 2026")
        doc.add_paragraph(f"original bullet {number * 2 + 1}", style="List Paragraph")
        doc.add_paragraph(f"original bullet {number * 2 + 2}", style="List Paragraph")
    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("SKILLS")
    for category in ("Network", "Security", "AI", "Vision", "Programming"):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{category}: ").bold = True
        paragraph.add_run("original items")
    doc.save(path)


def test_renderer_only_rewrites_fixed_project_and_skill_slots(tmp_path):
    template, output = tmp_path / "template.docx", tmp_path / "resume.docx"
    make_template(template)
    renderer.render_docx(
        template=template, output=output,
        project_bullets=[{"slot": 1, "text": "Built validated security lab evidence."}],
        skill_lines=[{"category": "Security", "items": "PKI/TLS, OpenSSL"}],
        project_subtitles=[{"slot": 1, "text": "JD-relevant verified subtitle"}],
    )
    paragraphs = Document(output).paragraphs
    texts = [paragraph.text for paragraph in paragraphs]
    assert texts[0] == "CANDIDATE NAME | github.com/ahpxna"
    assert texts[1] == "EDUCATION"
    assert texts[3] == "Project 1 — JD-relevant verified subtitle | GitHub\tJan 2026"
    assert texts[4] == "Built validated security lab evidence."
    assert texts[5] == ""
    assert texts[7] == ""
    assert texts[-5] == "Security: PKI/TLS, OpenSSL"
    assert texts[-1] == "Programming: original items"


def test_renderer_refuses_more_content_than_the_one_page_slot_budget(tmp_path):
    template = tmp_path / "template.docx"
    make_template(template)
    try:
        renderer.render_docx(template=template, output=tmp_path / "out.docx",
                             project_bullets=[{"slot": index + 1, "text": "x"} for index in range(13)], skill_lines=[])
    except renderer.ResumeTemplateError as exc:
        assert "12" in str(exc)
    else:
        raise AssertionError("Expected the fixed slot budget to reject extra bullets")


def test_renderer_requires_primary_before_secondary_and_enforces_each_budget(tmp_path):
    template = tmp_path / "template.docx"
    make_template(template)
    try:
        renderer.render_docx(
            template=template, output=tmp_path / "secondary.docx",
            project_bullets=[{"slot": 2, "text": "A secondary fact without the primary bullet."}],
            skill_lines=[],
        )
    except renderer.ResumeTemplateError as exc:
        assert "Secondary slot 2" in str(exc)
    else:
        raise AssertionError("Secondary project bullet must require its primary bullet")

    try:
        renderer.render_docx(
            template=template, output=tmp_path / "long.docx",
            project_bullets=[{"slot": 1, "text": "x" * 201}], skill_lines=[],
        )
    except renderer.ResumeTemplateError as exc:
        assert "200" in str(exc)
    else:
        raise AssertionError("Primary project bullet must fit its fixed two-line budget")
