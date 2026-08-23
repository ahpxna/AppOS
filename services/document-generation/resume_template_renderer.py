"""Render a verified resume into the fixed one-page Word/PDF template.

Only the twelve existing project-bullet slots, five existing skill rows, and
the audited subtitle between a fixed project name and its GitHub link may
change. The header identity, dates, links, education, experience,
certifications, page geometry, and every style remain from the supplied Word
template. The renderer refuses to silently shrink text or accept a multi-page
PDF.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import shutil
import subprocess
from typing import Any

from docx import Document
from services.common.resume_project_header_audit import HEADER_SLOTS, subtitle_bounds


PROJECTS_HEADING = "PROJECTS"
CERTIFICATIONS_HEADING = "CERTIFICATIONS"
SKILLS_HEADING = "SKILLS"
MAX_PROJECT_BULLETS = 12
MAX_SKILL_LINES = 5
PRIMARY_BULLET_MAX_CHARS = 200  # visual budget: two lines at the fixed 9.5 pt width
SECONDARY_BULLET_MAX_CHARS = 105  # visual budget: one line
# Top-origin coordinates and line capacity measured from the authoritative
# one-page PDF, not inferred from LibreOffice's incompatible pagination.
PROJECT_BULLET_SLOTS = [(276, 2), (299, 2), (338, 2), (364, 1), (392, 2), (417, 2),
                        (456, 2), (481, 1), (510, 2), (532, 2), (569, 1), (583, 2)]
SKILL_SLOTS = [(681, 2), (706, 1), (722, 1), (735, 1), (747, 1)]


class ResumeTemplateError(RuntimeError):
    """The fixed template cannot safely accommodate the requested content."""


def _paragraph_index(paragraphs, heading: str) -> int:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().upper() == heading:
            return index
    raise ResumeTemplateError(f"Template heading not found: {heading}")


def _replace_plain(paragraph, text: str) -> None:
    """Replace visible text while retaining the source paragraph and run style."""
    if not paragraph.runs:
        raise ResumeTemplateError("Template slot has no formatted run.")
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_skill(paragraph, category: str, items: str) -> None:
    """Keep the template's bold category run and normal item run intact."""
    if len(paragraph.runs) < 2:
        raise ResumeTemplateError("Skill slot does not contain category and item runs.")
    paragraph.runs[0].text = f"{category.strip()}: "
    paragraph.runs[1].text = items.strip()
    for run in paragraph.runs[2:]:
        run.text = ""


def _is_project_header(paragraph) -> bool:
    text = paragraph.text.casefold()
    return "—" in paragraph.text and "|" in paragraph.text and "github" in text


def _slot_indices(paragraphs) -> tuple[list[int], list[int]]:
    """Locate the only paragraphs the renderer is ever allowed to mutate."""
    project_start = _paragraph_index(paragraphs, PROJECTS_HEADING)
    certification_start = _paragraph_index(paragraphs, CERTIFICATIONS_HEADING)
    skills_start = _paragraph_index(paragraphs, SKILLS_HEADING)
    bullets = [
        index for index, paragraph in enumerate(paragraphs[project_start + 1:certification_start], project_start + 1)
        if paragraph.style.name == "List Paragraph" and not _is_project_header(paragraph)
    ]
    skills = list(range(skills_start + 1, skills_start + 1 + MAX_SKILL_LINES))
    if len(bullets) != MAX_PROJECT_BULLETS or len(skills) != MAX_SKILL_LINES:
        raise ResumeTemplateError("Template slot map changed; re-distill before rendering.")
    return bullets, skills


def _header_indices(paragraphs) -> list[int]:
    """Locate the six immutable project header paragraphs, in template order."""
    project_start = _paragraph_index(paragraphs, PROJECTS_HEADING)
    certification_start = _paragraph_index(paragraphs, CERTIFICATIONS_HEADING)
    headers = [
        index for index, paragraph in enumerate(paragraphs[project_start + 1:certification_start], project_start + 1)
        if _is_project_header(paragraph)
    ]
    if len(headers) != len(HEADER_SLOTS):
        raise ResumeTemplateError("Template project header map changed; re-distill before rendering.")
    return headers


def _header_shell(paragraph) -> tuple[str, str]:
    """Snapshot every header character except the explicitly editable subtitle."""
    start, end = subtitle_bounds(paragraph.text)
    return paragraph.text[:start], paragraph.text[end:]


def _replace_subtitle(paragraph, text: str) -> None:
    """Replace only the text after — and before | while keeping link/date runs."""
    regular = "".join(run.text for run in paragraph.runs)
    start, end = subtitle_bounds(regular)
    overlap = []
    cursor = 0
    for run in paragraph.runs:
        run_start, run_end = cursor, cursor + len(run.text)
        if run_end > start and run_start < end:
            overlap.append((run, run_start, run_end))
        cursor = run_end
    if not overlap:
        raise ResumeTemplateError("Project subtitle slot has no editable formatted run.")
    # Prefer the template's italic subtitle run. This preserves the intended
    # visual distinction even in ApplyOps, whose legacy subtitle spans runs.
    target = next((run for run, _, _ in overlap if run.italic), overlap[0][0])
    for run, run_start, run_end in overlap:
        # Preserve characters that share a run with the subtitle, e.g. a
        # simplified test template may keep name/subtitle/link in one run.
        original = run.text
        left = original[:max(0, start - run_start)] if run_start < start else ""
        right = original[max(0, end - run_start):] if run_end > end else ""
        run.text = left + (text if run is target else "") + right


def protected_snapshot(document) -> dict[str, Any]:
    """Capture all text and hyperlink targets outside permitted text slots."""
    paragraphs = document.paragraphs
    bullets, skills = _slot_indices(paragraphs)
    headers = _header_indices(paragraphs)
    editable = set(bullets + skills + headers)
    hyperlinks = sorted(
        rel.target_ref for rel in document.part.rels.values()
        if rel.reltype.endswith("/hyperlink")
    )
    return {
        "paragraphs": {index: paragraph.text for index, paragraph in enumerate(paragraphs) if index not in editable},
        "project_header_shells": {index: _header_shell(paragraphs[index]) for index in headers},
        "hyperlinks": hyperlinks,
    }


def render_docx(*, template: Path, output: Path, project_bullets: list[dict[str, Any]],
                skill_lines: list[dict[str, str]], project_subtitles: list[dict[str, Any]] | None = None) -> None:
    """Copy the template and alter only fixed bullet/skill slots in place."""
    if len(project_bullets) > MAX_PROJECT_BULLETS:
        raise ResumeTemplateError(f"Template has {MAX_PROJECT_BULLETS} project-bullet slots.")
    if len(skill_lines) > MAX_SKILL_LINES:
        raise ResumeTemplateError(f"Template has {MAX_SKILL_LINES} skill-category rows.")
    project_subtitles = project_subtitles or []
    if len(project_subtitles) > len(HEADER_SLOTS):
        raise ResumeTemplateError(f"Template has {len(HEADER_SLOTS)} editable project subtitles.")
    if not template.is_file():
        raise ResumeTemplateError(f"Resume template not found: {template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    document = Document(output)
    immutable = protected_snapshot(document)
    bullet_indices, skill_indices = _slot_indices(document.paragraphs)
    header_indices = _header_indices(document.paragraphs)
    bullet_slots = [document.paragraphs[index] for index in bullet_indices]
    skill_slots = [document.paragraphs[index] for index in skill_indices]
    header_slots = dict(zip(HEADER_SLOTS, (document.paragraphs[index] for index in header_indices)))
    seen_slots: set[int] = set()
    requested_slots: list[int] = []
    for item in project_bullets:
        try:
            requested_slots.append(int(item.get("slot")))
        except (TypeError, ValueError) as exc:
            raise ResumeTemplateError("Every project update needs an integer fixed slot 1..12.") from exc
    for slot_number in requested_slots:
        if not 1 <= slot_number <= MAX_PROJECT_BULLETS:
            raise ResumeTemplateError("Project updates must use unique fixed slots 1..12.")
        if slot_number % 2 == 0 and slot_number - 1 not in requested_slots:
            raise ResumeTemplateError(
                f"Secondary slot {slot_number} requires its matching primary slot {slot_number - 1}."
            )
    validated_updates: list[tuple[int, str]] = []
    for item in project_bullets:
        try:
            slot_number = int(item.get("slot"))
        except (TypeError, ValueError) as exc:
            raise ResumeTemplateError("Every project update needs an integer fixed slot 1..12.") from exc
        if not 1 <= slot_number <= MAX_PROJECT_BULLETS or slot_number in seen_slots:
            raise ResumeTemplateError("Project updates must use unique fixed slots 1..12.")
        seen_slots.add(slot_number)
        compact = " ".join(str(item.get("text") or "").split())
        limit = PRIMARY_BULLET_MAX_CHARS if slot_number % 2 else SECONDARY_BULLET_MAX_CHARS
        if not compact or len(compact) > limit:
            kind = "primary (two-line)" if slot_number % 2 else "secondary (one-line)"
            raise ResumeTemplateError(f"Slot {slot_number} is {kind}; its text must be 1..{limit} characters.")
        validated_updates.append((slot_number, compact))
    # A generated resume must not inherit irrelevant bullets from the generic
    # template.  Clearing text is permitted; headings, dates, links and list
    # paragraph formatting remain intact for every project block.
    for bullet in bullet_slots:
        _replace_plain(bullet, "")
    for slot_number, compact in validated_updates:
        _replace_plain(bullet_slots[slot_number - 1], compact)
    subtitle_slots: set[int] = set()
    for item in project_subtitles:
        try:
            slot_number = int(item.get("slot"))
        except (TypeError, ValueError) as exc:
            raise ResumeTemplateError("Every project subtitle update needs a fixed header slot.") from exc
        compact = " ".join(str(item.get("text") or "").split())
        if slot_number not in header_slots or slot_number in subtitle_slots:
            raise ResumeTemplateError("Project subtitle updates must use unique header slots 1, 3, 5, 7, 9, or 11.")
        if not compact or len(compact) > 88:
            raise ResumeTemplateError("Project subtitle must fit the fixed one-line 88-character budget.")
        subtitle_slots.add(slot_number)
        _replace_subtitle(header_slots[slot_number], compact)
    for slot, item in zip(skill_slots, skill_lines):
        category = str(item.get("category") or "").strip()
        values = str(item.get("items") or "").strip()
        if not category or not values or len(category) > 45 or len(values) > 220:
            raise ResumeTemplateError("Each tailored skill row needs a short category and items.")
        _replace_skill(slot, category, values)
    document.save(output)
    if protected_snapshot(Document(output)) != immutable:
        raise ResumeTemplateError(
            "Template integrity check failed: a protected name/contact/project-heading/hyperlink field changed."
        )


def export_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Export through LibreOffice; fail instead of pretending a wrong layout passed."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise ResumeTemplateError("LibreOffice is required for PDF export (install libreoffice).")
    output_dir.mkdir(parents=True, exist_ok=True)
    proc = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        text=True, capture_output=True, timeout=120,
    )
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if proc.returncode or not pdf_path.is_file():
        detail = (proc.stderr or proc.stdout).strip()[-500:]
        raise ResumeTemplateError(f"LibreOffice PDF export failed: {detail or 'no output'}")
    pdfinfo = shutil.which("pdfinfo")
    if not pdfinfo:
        raise ResumeTemplateError("Poppler pdfinfo is required for one-page validation (install poppler-utils).")
    info = subprocess.run([pdfinfo, str(pdf_path)], text=True, capture_output=True, timeout=30, check=True).stdout
    pages = next((line.split(":", 1)[1].strip() for line in info.splitlines() if line.startswith("Pages:")), "0")
    if pages != "1":
        raise ResumeTemplateError(
            f"PDF rendered as {pages} pages, not one. Do not shrink the font; shorten project/skill content or install Arial-compatible fonts."
        )
    return pdf_path


def _arial_font() -> Path:
    """Find a real Arial-compatible TTF instead of silently changing metrics."""
    configured = Path(__import__("os").getenv("JOBOS_RESUME_ARIAL_TTF", "")).expanduser()
    candidates = [configured] if str(configured) != "." else []
    candidates += [Path("/Library/Fonts/Arial.ttf"), Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf")]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    fc_match = shutil.which("fc-match")
    if fc_match:
        found = subprocess.run([fc_match, "-f", "%{file}", "Arial"], text=True, capture_output=True).stdout.strip()
        if found and Path(found).is_file() and "arial" in Path(found).name.lower():
            return Path(found)
    raise ResumeTemplateError("Arial TTF is required for pixel-stable PDF output. Set JOBOS_RESUME_ARIAL_TTF or install msttcorefonts.")


def export_pdf_from_reference(*, reference_pdf: Path, output_pdf: Path,
                              project_bullets: list[str], skill_lines: list[dict[str, str]]) -> Path:
    """Overlay only editable slots on the known-good one-page PDF reference."""
    if not reference_pdf.is_file():
        raise ResumeTemplateError(f"Reference PDF not found: {reference_pdf}")
    if len(project_bullets) > MAX_PROJECT_BULLETS or len(skill_lines) > MAX_SKILL_LINES:
        raise ResumeTemplateError("Tailoring exceeds the fixed one-page PDF slot budget.")
    try:
        from pypdf import PdfReader, PdfWriter
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ResumeTemplateError("Install document-generation requirements for PDF overlay export.") from exc
    font = _arial_font()
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    overlay = output_pdf.with_suffix(".overlay.pdf")
    pdfmetrics.registerFont(TTFont("JobOSArial", str(font)))
    page_height, right = 792, 568
    canvas_out = canvas.Canvas(str(overlay), pagesize=(612, page_height))
    canvas_out.setFillColorRGB(1, 1, 1)
    for top, lines in PROJECT_BULLET_SLOTS:
        canvas_out.rect(54, page_height - top - (lines * 11.2), 520, lines * 11.2, fill=1, stroke=0)
    for top, lines in SKILL_SLOTS:
        canvas_out.rect(40, page_height - top - (lines * 11.2), 535, lines * 11.2, fill=1, stroke=0)
    canvas_out.setFillColorRGB(0, 0, 0)
    canvas_out.setFont("JobOSArial", 9.5)
    for (top, capacity), text in zip(PROJECT_BULLET_SLOTS, project_bullets):
        words, lines, current = text.split(), [], ""
        for word in words:
            proposed = (current + " " + word).strip()
            if pdfmetrics.stringWidth(proposed, "JobOSArial", 9.5) > 492 and current:
                lines.append(current); current = word
            else: current = proposed
        if current: lines.append(current)
        if len(lines) > capacity:
            raise ResumeTemplateError("A tailored project bullet exceeds its original PDF line capacity.")
        canvas_out.drawString(58, page_height - top - 8, u"•")
        for index, line in enumerate(lines): canvas_out.drawString(72, page_height - top - 8 - index * 11.2, line)
    for (top, capacity), item in zip(SKILL_SLOTS, skill_lines):
        text = f"{item['category']}: {item['items']}"
        if pdfmetrics.stringWidth(text, "JobOSArial", 9.5) > 525 * capacity:
            raise ResumeTemplateError("A tailored skill row exceeds its original PDF line capacity.")
        canvas_out.drawString(43.6, page_height - top - 8, text)
    canvas_out.save()
    writer = PdfWriter(); base = PdfReader(str(reference_pdf)); layer = PdfReader(str(overlay))
    base.pages[0].merge_page(layer.pages[0]); writer.add_page(base.pages[0])
    with output_pdf.open("wb") as stream: writer.write(stream)
    overlay.unlink(missing_ok=True)
    return output_pdf
