"""
Parse data/profile_sources_v2/**/*.{pdf,docx} into plain-text .txt files
under data/profile_parsed_v2/, in the exact layout ingest_profile_sources_v2.py
expects: find_parsed_path() looks up `(PARSED_ROOT / rel_path).with_suffix(".txt")`
for every raw file under SOURCE_ROOT.

Why this script exists: ingest_profile_sources_v2.py does NOT parse
anything itself -- it only registers raw files in the DB and looks for an
already-parsed .txt counterpart. On a fresh machine, nothing has ever
populated data/profile_parsed_v2, so every file shows up as
`parsed: MISSING, sections: 0, chunks: 0` and every downstream step (doc
mapping, evidence extraction, chunk embedding, capability building, brief
generation, context packs) has nothing to work with and fails or no-ops.
Confirmed live 2026-08-01: `data/profile_parsed_v2` didn't even exist.

Two ways to fill data/profile_parsed_v2:
  1. build_profile_parsed_v2_from_existing.py (repo root) -- copies over
     already-parsed .txt files from the OLD data/profile_parsed/ directory
     for any data/profile_sources_v2 file that's byte-identical (same
     sha256) to a file that was already parsed under the old pipeline.
     Fast, no re-parsing, but only works for files carried over from the
     old layout.
  2. This script -- actually parses PDF/DOCX into text with pypdf /
     python-docx, for whatever's left. Slower, but works for genuinely
     new files.

Run (1) first, then (2) for whatever it didn't cover -- this script skips
any file that already has a parsed .txt (unless --force), so running both
in sequence is safe and non-redundant.

Also filters out Windows NTFS "Zone.Identifier" alternate-data-stream
sidecar files (e.g. `report.pdf:Zone.Identifier`) -- these are not real
documents, they're metadata WSL/Windows tooling sometimes copies over
literally as separate files, and ingest_profile_sources_v2.py does not
filter them out (plain `rglob("*")`), so without this filter they show up
as bogus "documents" with unclassified role and zero content, doubling
the apparent file count for no reason. Confirmed live: exactly half of a
reported 88 "files" were `:Zone.Identifier` companions of the other half.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

SOURCE_ROOT = Path("data/profile_sources_v2")
PARSED_ROOT = Path("data/profile_parsed_v2")

SUPPORTED_SUFFIXES = {".pdf", ".docx"}


def is_junk(path: Path) -> bool:
    # Windows "mark of the web" ADS sidecar, e.g. "foo.pdf:Zone.Identifier"
    return ":Zone.Identifier" in path.name or path.name.endswith("Zone.Identifier")


def parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"--- PAGE {i + 1} ---\n{text}")
    return "\n\n".join(parts)


def parse_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def main() -> int:
    ap = argparse.ArgumentParser(description="Parse profile_sources_v2 PDFs/DOCX into plain text.")
    ap.add_argument("--source-root", default=str(SOURCE_ROOT))
    ap.add_argument("--parsed-root", default=str(PARSED_ROOT))
    ap.add_argument("--force", action="store_true", help="Re-parse even if a .txt already exists.")
    ap.add_argument("--dry-run", action="store_true", help="List what would be parsed, write nothing.")
    args = ap.parse_args()

    source_root = Path(args.source_root)
    parsed_root = Path(args.parsed_root)
    parsed_root.mkdir(parents=True, exist_ok=True)

    if not source_root.exists():
        print(f"Source root not found: {source_root}", file=sys.stderr)
        return 1

    files = sorted(p for p in source_root.rglob("*") if p.is_file() and not p.name.startswith("."))

    skipped_junk = 0
    skipped_unsupported = 0
    skipped_existing = 0
    parsed_ok = 0
    parsed_fail = 0

    for f in files:
        if is_junk(f):
            skipped_junk += 1
            continue
        rel = f.relative_to(source_root)
        if f.suffix.lower() not in SUPPORTED_SUFFIXES:
            print(f"SKIP (unsupported type) {rel}")
            skipped_unsupported += 1
            continue

        out = (parsed_root / rel).with_suffix(".txt")
        if out.exists() and not args.force:
            skipped_existing += 1
            continue

        if args.dry_run:
            print(f"WOULD PARSE {rel} -> {out.relative_to(parsed_root)}")
            continue

        try:
            if f.suffix.lower() == ".pdf":
                text = parse_pdf(f)
            else:
                text = parse_docx(f)
        except Exception as e:
            print(f"FAIL {rel}: {e}")
            parsed_fail += 1
            continue

        if not text.strip():
            print(f"WARN {rel}: extracted empty text (scanned/image-only PDF? needs OCR, not this script)")

        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(text, encoding="utf-8")
        print(f"OK   {rel} -> {out.relative_to(parsed_root)}")
        parsed_ok += 1

    print("")
    print("===== SUMMARY =====")
    print(f"Total files under source root: {len(files)}")
    print(f"Zone.Identifier junk skipped:  {skipped_junk}")
    print(f"Unsupported type skipped:      {skipped_unsupported}")
    print(f"Already parsed, skipped:       {skipped_existing}")
    print(f"Parsed OK:                     {parsed_ok}")
    print(f"Parse failures:                {parsed_fail}")
    return 1 if parsed_fail else 0


if __name__ == "__main__":
    sys.exit(main())
