import hashlib
import json
import mimetypes
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

import psycopg
from psycopg.types.json import Jsonb

# PATCH H1: hard-fail. Thieu dependency thi chet ngay, khong parse rong im lang.
try:
    from pypdf import PdfReader
except ImportError as e:
    raise SystemExit(
        "FATAL: thieu 'pypdf'. Chay: pip install -r requirements.txt\n"
        f"       chi tiet: {e}"
    )

try:
    import docx
except ImportError as e:
    raise SystemExit(
        "FATAL: thieu 'python-docx'. Chay: pip install -r requirements.txt\n"
        f"       chi tiet: {e}"
    )


PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT))
from services.common.config import database_dsn

RAW_DIR = PROJECT_ROOT / "data" / "profile_raw"
PARSED_DIR = PROJECT_ROOT / "data" / "profile_parsed"
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".markdown",
}


@dataclass
class ParsedFile:
    path: Path
    sha256: str
    file_type: str
    mime_type: str
    text: str
    parser_used: str
    parse_status: str
    parse_error: Optional[str]


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def estimate_tokens(text: str) -> int:
    # Rough estimate. Good enough for chunk metadata.
    # English-like text usually ~4 chars/token.
    return max(1, len(text) // 4)


def parse_txt_or_md(path: Path) -> Tuple[str, str]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return normalize_text(raw), "plain_text"


def parse_pdf(path: Path) -> Tuple[str, str, Optional[str]]:
    try:
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception as e:
                page_text = f"\n[PAGE {i + 1} EXTRACTION ERROR: {e}]\n"
            if page_text.strip():
                pages.append(f"\n\n--- PAGE {i + 1} ---\n{page_text}")
        text = normalize_text("\n".join(pages))

        if len(text) < 200:
            # Could be scanned PDF or image-heavy PDF.
            return text, "pypdf", "LOW_TEXT_POSSIBLE_SCANNED_PDF"

        return text, "pypdf", None
    except Exception as e:
        return "", "pypdf", str(e)


def parse_docx(path: Path) -> Tuple[str, str, Optional[str]]:
    try:
        d = docx.Document(str(path))
        parts = []

        # PATCH H2: giu phan cap Heading tu Word.
        #   Heading N  -> "#### Title"  (MARKDOWN_HEADING_RE o section builder bat duoc)
        #   nhan in dam ket thuc bang ":" -> "@@FIELD@@ Nhan:" (KHONG phai heading)
        # Nho vay detect_heading khong con phai doan, va khong the nham
        # "Portfolio positioning:" thanh tieu de section.
        for para in d.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue

            style = (para.style.name or "").lower() if para.style is not None else ""

            if style.startswith("heading"):
                digits = "".join(c for c in style if c.isdigit())
                level = min(int(digits), 5) if digits else 1
                parts.append("")
                parts.append(f"{'#' * level} {txt}")
                parts.append("")
                continue

            if style in ("title", "subtitle"):
                parts.append("")
                parts.append(f"# {txt}")
                parts.append("")
                continue

            is_label = (
                len(txt) < 60
                and txt.endswith(":")
                and para.runs
                and bool(para.runs[0].bold)
            )
            if is_label:
                parts.append(f"@@FIELD@@ {txt}")
                continue

            parts.append(txt)

        # Preserve table rows roughly.
        for table_idx, table in enumerate(d.tables):
            parts.append(f"\n--- TABLE {table_idx + 1} ---")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return normalize_text("\n".join(parts)), "python-docx", None
    except Exception as e:
        return "", "python-docx", str(e)


def classify_file_type(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in [".txt"]:
        return "txt"
    if ext in [".md", ".markdown"]:
        return "markdown"
    return "unsupported"


def parse_file(path: Path) -> ParsedFile:
    sha = sha256_file(path)
    file_type = classify_file_type(path)
    mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"

    if file_type == "unsupported":
        return ParsedFile(
            path=path,
            sha256=sha,
            file_type=file_type,
            mime_type=mime_type,
            text="",
            parser_used="none",
            parse_status="unsupported",
            parse_error=f"Unsupported file extension: {path.suffix}",
        )

    if file_type in ["txt", "markdown"]:
        try:
            text, parser = parse_txt_or_md(path)
            return ParsedFile(path, sha, file_type, mime_type, text, parser, "parsed", None)
        except Exception as e:
            return ParsedFile(path, sha, file_type, mime_type, "", "plain_text", "parse_failed", str(e))

    if file_type == "pdf":
        text, parser, err = parse_pdf(path)
        if err == "LOW_TEXT_POSSIBLE_SCANNED_PDF":
            return ParsedFile(path, sha, file_type, mime_type, text, parser, "needs_ocr", err)
        if err:
            return ParsedFile(path, sha, file_type, mime_type, text, parser, "parse_failed", err)
        return ParsedFile(path, sha, file_type, mime_type, text, parser, "parsed", None)

    if file_type == "docx":
        text, parser, err = parse_docx(path)
        if err:
            return ParsedFile(path, sha, file_type, mime_type, text, parser, "parse_failed", err)
        return ParsedFile(path, sha, file_type, mime_type, text, parser, "parsed", None)

    return ParsedFile(path, sha, file_type, mime_type, "", "none", "unsupported", "Unsupported")



SECTION_KEYWORDS = {
    "education": "academic",
    "academic": "academic",
    "coursework": "academic",
    "transcript": "academic",
    "gpa": "academic",
    "degree": "academic",

    "skills": "skills",
    "technical skills": "skills",
    "tools": "skills",
    "technologies": "skills",

    "projects": "projects",
    "project": "projects",
    "portfolio": "projects",

    "experience": "experience",
    "professional experience": "experience",
    "work experience": "experience",
    "employment": "experience",
    "tutor": "experience",
    "intern": "experience",

    "research": "research",
    "independent study": "research",

    "certifications": "certifications",
    "certification": "certifications",

    "summary": "career_positioning",
    "profile": "career_positioning",
    "positioning": "career_positioning",
    "career": "career_positioning",
}


def clean_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"\s+", " ", line)
    return line


def looks_like_heading(line: str) -> bool:
    line = clean_line(line)
    if not line:
        return False

    # Short all-caps heading: EDUCATION, PROFESSIONAL EXPERIENCE
    if len(line) <= 80 and line.upper() == line and re.search(r"[A-Z]", line):
        return True

    # Common title-like headings
    low = line.lower().strip(":")
    if low in SECTION_KEYWORDS:
        return True

    # Lines that are short and don't end like a sentence often act as headings
    if len(line) <= 70 and not line.endswith((".", ",", ";")):
        if any(k in low for k in SECTION_KEYWORDS):
            return True

    return False


def category_from_text(text: str, filename: str = "") -> str:
    low = (filename + "\n" + text[:1200]).lower()

    # Prioritize explicit section headings in the chunk
    for key, category in SECTION_KEYWORDS.items():
        if re.search(rf"(^|\n)\s*{re.escape(key)}\s*(:|\n|$)", low):
            return category

    if any(k in low for k in ["coursework", "education", "degree", "university", "transcript", "gpa"]):
        return "academic"
    if any(k in low for k in ["python", "java", "sql", "linux", "networking", "cybersecurity", "nmap", "wireshark", "splunk"]):
        return "skills"
    if any(k in low for k in ["project", "built", "developed", "implemented", "designed"]):
        return "projects"
    if any(k in low for k in ["experience", "intern", "worked", "tutor", "employment", "professional"]):
        return "experience"
    if any(k in low for k in ["research", "marl", "causal", "experiment"]):
        return "research"

    return "unknown"


def split_blocks(text: str) -> List[str]:
    """
    Convert text into clean paragraph-ish blocks.
    Keeps headings as separate blocks when possible.
    """
    raw_lines = text.splitlines()
    blocks = []
    current = []

    def flush():
        nonlocal current
        if current:
            block = normalize_text("\n".join(current))
            if block:
                blocks.append(block)
            current = []

    for raw in raw_lines:
        line = clean_line(raw)

        if not line:
            flush()
            continue

        if looks_like_heading(line):
            flush()
            blocks.append(line)
            continue

        current.append(line)

    flush()
    return blocks


def chunk_text(text: str, target_chars: int = 3000, max_chars: int = 3800) -> List[str]:
    """
    Heading/paragraph-aware chunker.
    No character-overlap, because overlap caused chunks to start mid-word.
    Context continuity will be handled later by retrieval and neighboring chunks.
    """
    blocks = split_blocks(text)
    chunks = []
    current_blocks = []
    current_len = 0

    def flush_current():
        nonlocal current_blocks, current_len
        if current_blocks:
            chunk = normalize_text("\n\n".join(current_blocks))
            if chunk:
                chunks.append(chunk)
            current_blocks = []
            current_len = 0

    for block in blocks:
        blen = len(block)

        # If a block alone is too large, split by sentence-ish boundaries.
        if blen > max_chars:
            flush_current()
            pieces = split_large_block(block, target_chars)
            chunks.extend(pieces)
            continue

        # Start a new chunk at major headings if current chunk is already meaningful.
        if looks_like_heading(block) and current_len > target_chars * 0.45:
            flush_current()

        if current_len + blen + 2 > max_chars:
            flush_current()

        current_blocks.append(block)
        current_len += blen + 2

    flush_current()
    return [c for c in chunks if c.strip()]


def split_large_block(block: str, target_chars: int = 3000) -> List[str]:
    sentences = re.split(r"(?<=[.!?])\s+", block)
    pieces = []
    current = ""

    for s in sentences:
        if not current:
            current = s
            continue

        if len(current) + len(s) + 1 <= target_chars:
            current += " " + s
        else:
            pieces.append(current.strip())
            current = s

    if current.strip():
        pieces.append(current.strip())

    # Fallback if sentence split did not help
    final = []
    for p in pieces:
        if len(p) <= target_chars * 1.5:
            final.append(p)
        else:
            start = 0
            while start < len(p):
                final.append(p[start:start + target_chars].strip())
                start += target_chars

    return final


def infer_section(chunk: str, filename: str, idx: int) -> str:
    lines = [clean_line(l) for l in chunk.splitlines() if clean_line(l)]
    if not lines:
        return f"chunk_{idx}"

    first = lines[0]
    if looks_like_heading(first):
        return first[:90]

    low = chunk[:800].lower()

    if "professional experience" in low or "experience" in low:
        return "Experience"
    if "education" in low or "coursework" in low or "degree" in low:
        return "Education / Coursework"
    if "technical skills" in low or "skills" in low:
        return "Skills"
    if "project" in low:
        return "Projects"
    if "research" in low:
        return "Research"

    return f"{filename} chunk {idx}"
def write_parsed_text(parsed: ParsedFile) -> Optional[Path]:
    if not parsed.text:
        return None
    PARSED_DIR.mkdir(parents=True, exist_ok=True)
    out = PARSED_DIR / f"{parsed.path.stem}.{parsed.sha256[:10]}.txt"
    out.write_text(parsed.text, encoding="utf-8")
    return out


def upsert_raw_file(conn, parsed: ParsedFile, parsed_text_path: Optional[Path]) -> str:
    storage_url = f"local://{parsed.path.relative_to(PROJECT_ROOT)}"
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO raw_files (
              file_name,
              file_type,
              mime_type,
              storage_url,
              sha256,
              source,
              document_date,
              parse_status,
              parser_used,
              parse_error,
              is_active
            )
            VALUES (
              %s, %s, %s, %s, %s,
              'local_profile_ingestion',
              NULL,
              %s, %s, %s,
              true
            )
            ON CONFLICT (sha256)
            DO UPDATE SET
              file_name = EXCLUDED.file_name,
              file_type = EXCLUDED.file_type,
              mime_type = EXCLUDED.mime_type,
              storage_url = EXCLUDED.storage_url,
              parse_status = EXCLUDED.parse_status,
              parser_used = EXCLUDED.parser_used,
              parse_error = EXCLUDED.parse_error,
              is_active = true
            RETURNING id;
            """,
            (
                parsed.path.name,
                parsed.file_type,
                parsed.mime_type,
                storage_url,
                parsed.sha256,
                parsed.parse_status,
                parsed.parser_used,
                parsed.parse_error,
            ),
        )
        return str(cur.fetchone()[0])


def replace_chunks(conn, raw_file_id: str, parsed: ParsedFile) -> int:
    with conn.cursor() as cur:
        cur.execute("DELETE FROM profile_chunks WHERE file_id = %s;", (raw_file_id,))

    if parsed.parse_status not in ["parsed", "needs_ocr"]:
        return 0

    if not parsed.text.strip():
        return 0

    chunks = chunk_text(parsed.text)
    # Category is assigned per chunk below, not globally per file.

    with conn.cursor() as cur:
        for idx, chunk in enumerate(chunks, start=1):
            section = infer_section(chunk, parsed.path.name, idx)
            metadata = {
                "source_file": parsed.path.name,
                "sha256": parsed.sha256,
                "parser_used": parsed.parser_used,
                "parse_status": parsed.parse_status,
                "chunking": "paragraph_aware_v1",
            }

            cur.execute(
                """
                INSERT INTO profile_chunks (
                  file_id,
                  chunk_index,
                  section,
                  category,
                  text_content,
                  page_number,
                  token_count,
                  metadata,
                  embedding
                )
                VALUES (
                  %s, %s, %s, %s, %s,
                  NULL,
                  %s,
                  %s,
                  NULL
                );
                """,
                (
                    raw_file_id,
                    idx,
                    section,
                    category_from_text(chunk, parsed.path.name),
                    chunk,
                    estimate_tokens(chunk),
                    Jsonb(metadata),
                ),
            )

    return len(chunks)


def infer_section(chunk: str, filename: str, idx: int) -> str:
    lines = [l.strip() for l in chunk.splitlines() if l.strip()]
    if not lines:
        return f"chunk_{idx}"

    first = lines[0]
    if len(first) <= 90 and not first.endswith("."):
        return first[:90]

    # keyword-based fallback
    low = chunk[:500].lower()
    if "education" in low or "coursework" in low:
        return "Education / Coursework"
    if "skill" in low:
        return "Skills"
    if "project" in low:
        return "Projects"
    if "experience" in low:
        return "Experience"

    return f"{filename} chunk {idx}"


def scan_files() -> List[Path]:
    RAW_DIR.mkdir(parents=True, exist_ok=True)

    files = []
    for path in RAW_DIR.rglob("*"):
        if path.is_file() and not path.name.startswith("."):
            if path.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(path)
            else:
                print(f"[SKIP unsupported extension] {path}")

    return sorted(files)


def main() -> int:
    files = scan_files()

    print("===== PROFILE INGESTION =====")
    print(f"Project root: {PROJECT_ROOT}")
    print(f"Raw dir:      {RAW_DIR}")
    print(f"Parsed dir:   {PARSED_DIR}")
    print(f"Files found:  {len(files)}")

    if not files:
        print("")
        print("No files found.")
        print(f"Put PDF/DOCX/TXT/MD files into: {RAW_DIR}")
        return 0

    summary = []

    with psycopg.connect(database_dsn(), autocommit=False) as conn:
        for path in files:
            print("")
            print(f"--- Ingesting: {path.name} ---")
            parsed = parse_file(path)
            parsed_text_path = write_parsed_text(parsed)

            raw_file_id = upsert_raw_file(conn, parsed, parsed_text_path)
            chunk_count = replace_chunks(conn, raw_file_id, parsed)

            conn.commit()

            text_chars = len(parsed.text or "")
            print(f"raw_file_id:  {raw_file_id}")
            print(f"file_type:    {parsed.file_type}")
            print(f"status:       {parsed.parse_status}")
            print(f"parser:       {parsed.parser_used}")
            print(f"text_chars:   {text_chars}")
            print(f"chunks:       {chunk_count}")
            if parsed.parse_error:
                print(f"parse_error:  {parsed.parse_error}")
            if parsed_text_path:
                print(f"parsed_text:  {parsed_text_path.relative_to(PROJECT_ROOT)}")

            summary.append(
                {
                    "file": path.name,
                    "raw_file_id": raw_file_id,
                    "status": parsed.parse_status,
                    "parser": parsed.parser_used,
                    "text_chars": text_chars,
                    "chunks": chunk_count,
                    "parse_error": parsed.parse_error,
                }
            )

    print("")
    print("===== SUMMARY =====")
    print(json.dumps(summary, indent=2, ensure_ascii=False))
    print("")
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
